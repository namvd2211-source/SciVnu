from __future__ import annotations

import json
import os
import html
import mimetypes
import re
import io
import threading
import time
import traceback
import urllib.parse
import uuid
import unicodedata
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional

import requests
from fastapi import FastAPI, File, HTTPException, Request, Response, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from config.release_config import current_version

if __package__ == "backend":
    from backend.backend_core import (
        arxiv_search,
        build_reference_section,
        clamp_int,
        cliproxy_chat_completions_url,
        core_search,
        CORE_API_KEY,
        DEFAULT_LLM_PROVIDER,
        DEFAULT_GEMINI_MODEL,
        estimated_reference_words,
        GOOGLE_GEMINI_OAUTH_SCOPES,
        LLMConfig,
        NODE_ORDER,
        WorkflowHooks,
        build_llm_config,
        count_words,
        default_model_for_provider,
        execute_workflow,
        guess_language_label,
        infer_requested_output_language,
        cliproxy_available,
        google_auth_file_save,
        google_credentials_from_auth_file,
        google_oauth_userinfo,
        get_provider_models,
        init_statuses,
        llm_complete_text,
        llm_config_for_role,
        llm_stream_text,
        localize_in_text_citations,
        normalize_for_match,
        normalize_language_label,
        openalex_search,
        read_json_file,
        resolve_vertex_runtime_credentials,
        SCOPUS_API_KEY,
        scopus_search,
        UNPAYWALL_EMAIL,
        unpaywall_lookup,
        write_json_file,
    )
else:
    from backend_core import (
        arxiv_search,
        build_reference_section,
        clamp_int,
        cliproxy_chat_completions_url,
        core_search,
        CORE_API_KEY,
        DEFAULT_LLM_PROVIDER,
        DEFAULT_GEMINI_MODEL,
        estimated_reference_words,
        GOOGLE_GEMINI_OAUTH_SCOPES,
        LLMConfig,
        NODE_ORDER,
        WorkflowHooks,
        build_llm_config,
        count_words,
        default_model_for_provider,
        execute_workflow,
        guess_language_label,
        infer_requested_output_language,
        cliproxy_available,
        google_auth_file_save,
        google_credentials_from_auth_file,
        google_oauth_userinfo,
        get_provider_models,
        init_statuses,
        llm_complete_text,
        llm_config_for_role,
        llm_stream_text,
        localize_in_text_citations,
        normalize_for_match,
        normalize_language_label,
        openalex_search,
        read_json_file,
        resolve_vertex_runtime_credentials,
        SCOPUS_API_KEY,
        scopus_search,
        UNPAYWALL_EMAIL,
        unpaywall_lookup,
        write_json_file,
    )

WEB_DIR = str(REPO_ROOT / "web")


class ChatMessage(BaseModel):
    role: Literal["user", "assistant", "system"] = "user"
    content: str = Field(default="", max_length=20000)


class WorkflowRequest(BaseModel):
    topic: str = Field(min_length=1, max_length=4000)
    language: Literal["English", "Vietnamese"] = "English"
    provider: Literal["Gemini"] = DEFAULT_LLM_PROVIDER
    model: Optional[str] = None
    target_word_count: Optional[int] = Field(default=None, ge=1000, le=20000)
    reference_target: Optional[int] = Field(default=None, ge=4, le=200)
    chat_history: List[ChatMessage] = Field(default_factory=list)
    attachment_ids: List[str] = Field(default_factory=list)
    search_filters: Dict[str, Any] = Field(default_factory=dict)


class ChatTurnRequest(BaseModel):
    messages: List[ChatMessage] = Field(default_factory=list)
    language: Literal["English", "Vietnamese"] = "English"
    provider: Literal["Gemini"] = DEFAULT_LLM_PROVIDER
    model: Optional[str] = None
    target_word_count: Optional[int] = Field(default=None, ge=1000, le=20000)
    reference_target: Optional[int] = Field(default=None, ge=4, le=200)
    attachment_ids: List[str] = Field(default_factory=list)
    search_filters: Dict[str, Any] = Field(default_factory=dict)


class DocxRenderRequest(BaseModel):
    markdown: str = Field(default="", max_length=300000)
    filename: str = Field(default="research-workflow.docx", max_length=200)


class ResourceSearchRequest(BaseModel):
    source: Literal["scopus", "core", "openalex", "arxiv"]
    topic: str = Field(default="", max_length=4000)
    max_results: int = Field(default=25, ge=1, le=100)


class ResourceUnpaywallRequest(BaseModel):
    doi: str = Field(default="", max_length=300)


class JobSnapshot(BaseModel):
    id: str
    status: Literal["queued", "running", "cancelling", "completed", "error", "cancelled"]
    created_at: float
    updated_at: float
    request: Dict[str, Any]
    node_statuses: Dict[str, str]
    logs: List[str]
    manager_output: str = ""
    outline: str = ""
    papers: List[Dict[str, Any]] = Field(default_factory=list)
    review_feedback: str = ""
    qa_summary: str = ""
    draft: str = ""
    final_markdown: str = ""
    actual_word_count: int = 0
    error: str = ""
    attachments: List[Dict[str, Any]] = Field(default_factory=list)


app = FastAPI(title="AI Agentic Research Backend", version=current_version())
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def add_local_network_headers(request: Request, call_next: Any) -> Response:
    private_network_preflight = (
        request.method.upper() == "OPTIONS"
        and request.headers.get("access-control-request-private-network", "").strip().lower() == "true"
        and bool(request.headers.get("origin"))
    )
    if private_network_preflight:
        requested_headers = request.headers.get("access-control-request-headers", "").strip()
        requested_method = request.headers.get("access-control-request-method", "").strip() or "GET"
        origin = request.headers.get("origin", "*").strip() or "*"
        response = Response(status_code=204)
        response.headers["Vary"] = "Origin"
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Credentials"] = "true"
        response.headers["Access-Control-Allow-Methods"] = requested_method
        response.headers["Access-Control-Allow-Headers"] = requested_headers or "*"
        response.headers["Access-Control-Allow-Private-Network"] = "true"
        response.headers["Access-Control-Max-Age"] = "600"
        response.headers["Private-Network-Access-Name"] = "AI Agentic Research Local Companion"
        response.headers["Private-Network-Access-ID"] = "research-local-companion"
        return response

    response = await call_next(request)
    response.headers["Access-Control-Allow-Private-Network"] = "true"
    response.headers["Private-Network-Access-Name"] = "AI Agentic Research Local Companion"
    response.headers["Private-Network-Access-ID"] = "research-local-companion"
    return response

_jobs: Dict[str, JobSnapshot] = {}
_jobs_lock = threading.Lock()
_job_stop_events: Dict[str, threading.Event] = {}
_job_stop_lock = threading.Lock()
AUTH_STORE_DIR = os.getenv(
    "AUTH_STORE_DIR",
    "/tmp/research-auth" if os.getenv("K_SERVICE") else os.path.join(os.getcwd(), ".auth"),
)
UPLOAD_STORE_DIR = os.path.join(AUTH_STORE_DIR, "uploads")
FRONTEND_ORIGIN = os.getenv("FRONTEND_ORIGIN", "https://sci-vnucea.web.app").strip()
PUBLIC_BACKEND_URL = os.getenv(
    "PUBLIC_BACKEND_URL",
    "https://research-backend-213473562655.asia-southeast1.run.app",
).strip().rstrip("/")
GOOGLE_OAUTH_CLIENT_ID = os.getenv("GOOGLE_OAUTH_CLIENT_ID", "").strip()
GOOGLE_OAUTH_CLIENT_SECRET = os.getenv("GOOGLE_OAUTH_CLIENT_SECRET", "").strip()
GOOGLE_OAUTH_PROJECT_ID = os.getenv("GOOGLE_OAUTH_PROJECT_ID", "").strip()
OAUTH_STORE_BACKEND = os.getenv(
    "OAUTH_STORE_BACKEND",
    "firestore" if os.getenv("K_SERVICE") else "file",
).strip().lower()
OAUTH_FIRESTORE_COLLECTION_PREFIX = os.getenv("OAUTH_FIRESTORE_COLLECTION_PREFIX", "research_oauth").strip()
_firestore_client_cache: Any = None
_firestore_client_lock = threading.Lock()
GEMINI_CLI_QUOTA_URL = "https://cloudcode-pa.googleapis.com/v1internal:retrieveUserQuota"
GEMINI_CLI_LOAD_CODE_ASSIST_URL = "https://cloudcode-pa.googleapis.com/v1internal:loadCodeAssist"
GEMINI_CLI_QUOTA_HEADERS = {
    "Content-Type": "application/json",
}
GEMINI_CLI_PREMIUM_TIER_IDS = {"g1-ultra-tier"}
GEMINI_CLI_CREDIT_TYPE = "GOOGLE_ONE_AI"
GEMINI_CLI_BUCKET_GROUPS = [
    {
        "id": "gemini-flash-lite-series",
        "label": "Gemini Flash Lite Series",
        "preferred_model_id": "gemini-2.5-flash-lite",
        "model_ids": ["gemini-2.5-flash-lite"],
    },
    {
        "id": "gemini-flash-series",
        "label": "Gemini Flash Series",
        "preferred_model_id": "gemini-3-flash-preview",
        "model_ids": ["gemini-3-flash-preview", "gemini-2.5-flash"],
    },
    {
        "id": "gemini-pro-series",
        "label": "Gemini Pro Series",
        "preferred_model_id": "gemini-3.1-pro-preview",
        "model_ids": ["gemini-3.1-pro-preview", "gemini-3-pro-preview", "gemini-2.5-pro"],
    },
]
GEMINI_CLI_TIER_LABELS = {
    "free-tier": "Free",
    "legacy-tier": "Legacy",
    "standard-tier": "Standard",
    "g1-pro-tier": "Pro",
    "g1-ultra-tier": "Ultra",
}


class JobCancelledError(RuntimeError):
    pass


def ensure_auth_store_dir() -> str:
    os.makedirs(AUTH_STORE_DIR, exist_ok=True)
    return AUTH_STORE_DIR


def ensure_upload_store_dir() -> str:
    os.makedirs(UPLOAD_STORE_DIR, exist_ok=True)
    return UPLOAD_STORE_DIR


def _upload_blob_path(upload_id: str, original_name: str) -> str:
    safe_name = re.sub(r"[^a-zA-Z0-9._-]+", "_", original_name or "upload").strip("._") or "upload"
    return os.path.join(ensure_upload_store_dir(), f"{upload_id}_{safe_name}")


def _upload_meta_path(upload_id: str) -> str:
    return os.path.join(ensure_upload_store_dir(), f"{upload_id}.json")


def _truncate_chars(text: str, limit: int) -> str:
    text = str(text or "")
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 1)].rstrip() + "…"


def _guess_mime_type(filename: str, fallback: str = "application/octet-stream") -> str:
    guessed, _encoding = mimetypes.guess_type(filename or "")
    return guessed or fallback


def _extract_pdf_text_from_bytes(data: bytes, *, max_pages: int = 20, max_chars: int = 24000) -> str:
    try:
        import fitz
    except Exception:
        return ""
    parts: List[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for index in range(min(doc.page_count, max_pages)):
            parts.append(doc.load_page(index).get_text("text"))
            if sum(len(part) for part in parts) >= max_chars * 2:
                break
    return _truncate_chars("\n".join(parts).strip(), max_chars)


def _extract_docx_text_from_bytes(data: bytes, *, max_chars: int = 24000) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    import io

    document = Document(io.BytesIO(data))
    parts: List[str] = []
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        if text:
            parts.append(text)
        if sum(len(part) for part in parts) >= max_chars * 2:
            break
    for table in document.tables:
        row_lines: List[str] = []
        for row in table.rows[:20]:
            row_lines.append(" | ".join(str(cell.text or "").strip() for cell in row.cells[:8]))
        if row_lines:
            parts.append("\n".join(row_lines))
        if sum(len(part) for part in parts) >= max_chars * 2:
            break
    return _truncate_chars("\n\n".join(parts).strip(), max_chars)


def _extract_spreadsheet_text_from_bytes(data: bytes, *, max_chars: int = 20000) -> str:
    try:
        from openpyxl import load_workbook
    except Exception:
        return ""
    import io

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: List[str] = []
    for sheet in workbook.worksheets[:6]:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(min_row=1, max_row=60, values_only=True):
            cells = [str(cell).strip() for cell in row[:12] if cell not in (None, "")]
            if cells:
                parts.append("\t".join(cells))
            if sum(len(part) for part in parts) >= max_chars * 2:
                break
        if sum(len(part) for part in parts) >= max_chars * 2:
            break
    return _truncate_chars("\n".join(parts).strip(), max_chars)


def _extract_plain_text_from_bytes(data: bytes, *, max_chars: int = 20000) -> str:
    for encoding in ["utf-8", "utf-8-sig", "cp1252", "latin-1"]:
        try:
            return _truncate_chars(data.decode(encoding), max_chars)
        except Exception:
            continue
    return ""


def _extract_attachment_text(filename: str, mime_type: str, data: bytes) -> str:
    lower_name = str(filename or "").lower()
    lower_mime = str(mime_type or "").lower()
    if lower_name.endswith(".pdf") or lower_mime == "application/pdf":
        return _extract_pdf_text_from_bytes(data)
    if lower_name.endswith(".docx") or "wordprocessingml.document" in lower_mime:
        return _extract_docx_text_from_bytes(data)
    if lower_name.endswith((".xlsx", ".xlsm")) or "spreadsheetml" in lower_mime:
        return _extract_spreadsheet_text_from_bytes(data)
    if lower_name.endswith((".txt", ".md", ".csv", ".json", ".tsv")) or lower_mime.startswith("text/"):
        return _extract_plain_text_from_bytes(data)
    return ""


def _upload_record(upload_id: str) -> Optional[Dict[str, Any]]:
    path = _upload_meta_path(upload_id)
    if not os.path.exists(path):
        return None
    try:
        data = read_json_file(path)
    except Exception:
        return None
    return data if isinstance(data, dict) else None


def _write_upload_record(upload_id: str, payload: Dict[str, Any]) -> None:
    write_json_file(_upload_meta_path(upload_id), payload)


def _web_oauth_state_path() -> str:
    return os.path.join(ensure_auth_store_dir(), "web-google-oauth-state.json")


def _web_oauth_sessions_path() -> str:
    return os.path.join(ensure_auth_store_dir(), "web-google-oauth-sessions.json")


def _local_cli_proxy_auth_state_path() -> str:
    return os.path.join(ensure_auth_store_dir(), "local-cli-proxy-auth-state.json")


def _read_store(path: str, default: Dict[str, Any]) -> Dict[str, Any]:
    if not os.path.exists(path):
        return dict(default)
    try:
        data = read_json_file(path)
    except Exception:
        return dict(default)
    merged = dict(default)
    merged.update(data)
    return merged


def _write_store(path: str, data: Dict[str, Any]) -> None:
    write_json_file(path, data)


def _read_web_oauth_state() -> Dict[str, Any]:
    return _read_store(_web_oauth_state_path(), {"flows": {}})


def _write_web_oauth_state(data: Dict[str, Any]) -> None:
    _write_store(_web_oauth_state_path(), data)


def _read_web_oauth_sessions() -> Dict[str, Any]:
    return _read_store(_web_oauth_sessions_path(), {"sessions": {}})


def _write_web_oauth_sessions(data: Dict[str, Any]) -> None:
    _write_store(_web_oauth_sessions_path(), data)


def _read_local_cli_proxy_auth_state() -> Dict[str, Any]:
    return _read_store(
        _local_cli_proxy_auth_state_path(),
        {"status": "unknown", "reason": "", "updated_at": 0.0},
    )


def _write_local_cli_proxy_auth_state(data: Dict[str, Any]) -> None:
    _write_store(_local_cli_proxy_auth_state_path(), data)


def _mark_local_cli_proxy_auth_state(status: str, reason: str = "") -> None:
    _write_local_cli_proxy_auth_state(
        {
            "status": str(status or "unknown").strip() or "unknown",
            "reason": str(reason or "").strip(),
            "updated_at": time.time(),
        }
    )


def _clear_local_cli_proxy_auth_state() -> None:
    _mark_local_cli_proxy_auth_state("valid", "")


def _is_local_cli_proxy_auth_error_message(message: str) -> bool:
    lowered = str(message or "").strip().lower()
    return any(
        token in lowered
        for token in [
            "401",
            "unauthenticated",
            "invalid authentication credentials",
            "expected oauth 2 access token",
            "local gemini auth token is missing or expired",
            "sign in again",
            "access token",
            "login cookie",
        ]
    )


def _is_local_companion_runtime() -> bool:
    return PUBLIC_BACKEND_URL.startswith("http://127.0.0.1") or PUBLIC_BACKEND_URL.startswith("http://localhost")


def _local_cli_proxy_auth_dir() -> Path:
    configured = os.getenv("RESEARCH_COMPANION_CLI_PROXY_AUTH_DIR", "").strip()
    if configured:
        return Path(configured)
    try:
        home = Path.home()
    except Exception:
        home = None
    if home:
        return home / ".cli-proxy-api"
    local_appdata = os.getenv("LOCALAPPDATA", "").strip()
    if local_appdata:
        return Path(local_appdata) / "ResearchCompanion" / "cli-proxy-auth"
    return Path.cwd() / ".research-companion" / "cli-proxy-auth"


def _local_cli_proxy_auth_dirs() -> List[Path]:
    dirs: List[Path] = []
    configured = os.getenv("RESEARCH_COMPANION_CLI_PROXY_AUTH_DIR", "").strip()
    if configured:
        dirs.append(Path(configured))
    try:
        home = Path.home()
    except Exception:
        home = None
    if home:
        dirs.append(home / ".cli-proxy-api")
    legacy_configured = os.getenv("RESEARCH_COMPANION_CLI_PROXY_AUTH_DIR_LEGACY", "").strip()
    if legacy_configured:
        dirs.append(Path(legacy_configured))
    local_appdata = os.getenv("LOCALAPPDATA", "").strip()
    if local_appdata:
        dirs.append(Path(local_appdata) / "ResearchCompanion" / "cli-proxy-auth")

    ordered: List[Path] = []
    seen: set[str] = set()
    for path in dirs:
        key = str(path.resolve()) if path.exists() else str(path)
        if key not in seen:
            ordered.append(path)
            seen.add(key)
    return ordered


def _value_as_text(value: Any) -> str:
    if isinstance(value, str):
        text = value.strip()
        return text
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return str(value)
    return ""


def _value_as_number(value: Any) -> Optional[float]:
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return float(value)
    text = _value_as_text(value)
    if not text:
        return None
    try:
        return float(text)
    except ValueError:
        return None


def _value_as_fraction(value: Any) -> Optional[float]:
    number = _value_as_number(value)
    if number is not None:
        return number
    text = _value_as_text(value)
    if text.endswith("%"):
        try:
            return float(text[:-1]) / 100.0
        except ValueError:
            return None
    return None


def _normalize_model_id(value: Any) -> str:
    model_id = _value_as_text(value)
    if model_id.endswith("_vertex"):
        model_id = model_id[: -len("_vertex")]
    return model_id


def _parse_datetime(value: Any) -> Optional[datetime]:
    text = _value_as_text(value)
    if not text:
        return None
    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00"))
    except ValueError:
        return None


def _load_local_cli_proxy_auth_entries() -> List[Dict[str, Any]]:
    entries: List[Dict[str, Any]] = []
    seen_paths: set[str] = set()
    for auth_dir in _local_cli_proxy_auth_dirs():
        if not auth_dir.exists():
            continue
        for path in auth_dir.glob("*.json"):
            try:
                resolved = str(path.resolve())
            except Exception:
                resolved = str(path)
            if resolved in seen_paths:
                continue
            seen_paths.add(resolved)
            try:
                raw = read_json_file(str(path))
            except Exception:
                continue
            auth_type = _value_as_text(raw.get("type") or raw.get("provider"))
            if auth_type.lower() not in {"gemini", "gemini-cli"}:
                continue
            if bool(raw.get("disabled")):
                continue
            token_info = raw.get("token") if isinstance(raw.get("token"), dict) else {}
            email = _value_as_text(raw.get("email"))
            project_id = _value_as_text(raw.get("project_id"))
            checked = bool(raw.get("checked", True))
            entry = {
                "name": path.name,
                "path": str(path),
                "email": email,
                "project_id": project_id,
                "checked": checked,
                "token": token_info,
                "mtime": path.stat().st_mtime,
            }
            entries.append(entry)
    entries.sort(key=lambda item: (not bool(item.get("checked")), -float(item.get("mtime") or 0.0), item.get("name", "")))
    return entries


def _persist_local_cli_proxy_token(entry: Dict[str, Any], token_info: Dict[str, Any]) -> None:
    path = _value_as_text(entry.get("path"))
    if not path:
        return
    try:
        payload = read_json_file(path)
    except Exception:
        payload = {}
    payload["token"] = token_info
    write_json_file(path, payload)


def _local_cli_proxy_token_expires_soon(token_info: Dict[str, Any], *, within_seconds: int = 300) -> bool:
    if not isinstance(token_info, dict):
        return True
    expiry = _parse_datetime(token_info.get("expiry"))
    if expiry is None:
        return True
    return expiry.timestamp() <= time.time() + max(30, int(within_seconds))


def _refresh_local_cli_proxy_token(entry: Dict[str, Any], *, force: bool = False) -> Optional[Dict[str, Any]]:
    token_info = entry.get("token") if isinstance(entry.get("token"), dict) else {}
    if not isinstance(token_info, dict):
        return None
    refresh_token = _value_as_text(token_info.get("refresh_token"))
    if not refresh_token:
        return None
    try:
        from google.auth.transport.requests import Request as GoogleAuthRequest
        from google.oauth2.credentials import Credentials

        credentials = Credentials(
            token=_value_as_text(token_info.get("access_token") or token_info.get("token")) or None,
            refresh_token=refresh_token or None,
            token_uri=_value_as_text(token_info.get("token_uri")) or "https://oauth2.googleapis.com/token",
            client_id=_value_as_text(token_info.get("client_id")) or None,
            client_secret=_value_as_text(token_info.get("client_secret")) or None,
            scopes=token_info.get("scopes") if isinstance(token_info.get("scopes"), list) else GOOGLE_GEMINI_OAUTH_SCOPES,
        )
        expiry = _parse_datetime(token_info.get("expiry"))
        if expiry is not None:
            credentials.expiry = expiry
        if force or credentials.expired or _local_cli_proxy_token_expires_soon(token_info):
            credentials.refresh(GoogleAuthRequest())
            refreshed_token = json.loads(credentials.to_json())
            if not _value_as_text(refreshed_token.get("refresh_token")):
                refreshed_token["refresh_token"] = refresh_token
            _persist_local_cli_proxy_token(entry, refreshed_token)
            entry["token"] = refreshed_token
            return refreshed_token
    except Exception:
        return None
    return token_info


def _cli_proxy_credentials_from_token_info(token_info: Dict[str, Any]) -> Any:
    if not isinstance(token_info, dict):
        return None
    try:
        from google.auth.transport.requests import Request as GoogleAuthRequest
        from google.oauth2.credentials import Credentials

        credentials = Credentials(
            token=_value_as_text(token_info.get("access_token") or token_info.get("token")) or None,
            refresh_token=_value_as_text(token_info.get("refresh_token")) or None,
            token_uri=_value_as_text(token_info.get("token_uri")) or "https://oauth2.googleapis.com/token",
            client_id=_value_as_text(token_info.get("client_id")) or None,
            client_secret=_value_as_text(token_info.get("client_secret")) or None,
            scopes=token_info.get("scopes") if isinstance(token_info.get("scopes"), list) else GOOGLE_GEMINI_OAUTH_SCOPES,
        )
        expiry = _parse_datetime(token_info.get("expiry"))
        if expiry is not None:
            credentials.expiry = expiry
        if credentials.expired and credentials.refresh_token:
            credentials.refresh(GoogleAuthRequest())
        return credentials
    except Exception:
        return None


def _resolve_local_cli_proxy_access_token(entry: Dict[str, Any]) -> str:
    token_info = entry.get("token") if isinstance(entry.get("token"), dict) else {}
    refreshed_token = _refresh_local_cli_proxy_token(entry)
    if isinstance(refreshed_token, dict):
        token_info = refreshed_token
    credentials = _cli_proxy_credentials_from_token_info(token_info)
    if credentials is not None and getattr(credentials, "token", None):
        refreshed_token = dict(token_info)
        refreshed_token["access_token"] = credentials.token
        try:
            refreshed_token["expiry"] = credentials.expiry.isoformat() if getattr(credentials, "expiry", None) else token_info.get("expiry")
        except Exception:
            pass
        _persist_local_cli_proxy_token(entry, refreshed_token)
        return str(credentials.token)
    return _value_as_text(token_info.get("access_token") or token_info.get("token"))


def _cloudcode_json_request(url: str, access_token: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    response = requests.post(
        url,
        headers={
            **GEMINI_CLI_QUOTA_HEADERS,
            "Authorization": f"Bearer {access_token}",
        },
        json=payload,
        timeout=25,
    )
    if response.status_code < 200 or response.status_code >= 300:
        detail = response.text.strip() or response.reason
        raise HTTPException(status_code=response.status_code, detail=detail)
    data = response.json()
    return data if isinstance(data, dict) else {}


def _local_cli_proxy_json_request(entry: Dict[str, Any], url: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    access_token = _resolve_local_cli_proxy_access_token(entry)
    if not access_token:
        raise HTTPException(
            status_code=401,
            detail="Local Gemini auth token is missing or expired. Sign in again in ResearchCompanion.exe.",
        )
    try:
        return _cloudcode_json_request(url, access_token, payload)
    except HTTPException as exc:
        if exc.status_code != 401:
            raise
        refreshed_token = _refresh_local_cli_proxy_token(entry, force=True)
        refreshed_access_token = _value_as_text(
            (refreshed_token or {}).get("access_token") or (refreshed_token or {}).get("token")
        )
        if not refreshed_access_token or refreshed_access_token == access_token:
            raise
        return _cloudcode_json_request(url, refreshed_access_token, payload)


def _gemini_cli_bucket_groups(raw_buckets: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    grouped: List[Dict[str, Any]] = []
    raw_by_model: Dict[str, Dict[str, Any]] = {}
    for bucket in raw_buckets:
        model_id = _normalize_model_id(bucket.get("modelId") or bucket.get("model_id"))
        if not model_id:
            continue
        raw_by_model[model_id] = {
            "model_id": model_id,
            "token_type": _value_as_text(bucket.get("tokenType") or bucket.get("token_type")),
            "remaining_fraction": _value_as_fraction(bucket.get("remainingFraction") or bucket.get("remaining_fraction")),
            "remaining_amount": _value_as_number(bucket.get("remainingAmount") or bucket.get("remaining_amount")),
            "reset_time": _value_as_text(bucket.get("resetTime") or bucket.get("reset_time")) or "",
        }

    for group in GEMINI_CLI_BUCKET_GROUPS:
        buckets = [raw_by_model[model_id] for model_id in group["model_ids"] if model_id in raw_by_model]
        if not buckets:
            continue
        preferred = next((item for item in buckets if item["model_id"] == group["preferred_model_id"]), buckets[0])
        fraction_values = [item["remaining_fraction"] for item in buckets if item["remaining_fraction"] is not None]
        amount_values = [item["remaining_amount"] for item in buckets if item["remaining_amount"] is not None]
        reset_values = [item["reset_time"] for item in buckets if item["reset_time"]]
        grouped.append(
            {
                "id": group["id"],
                "label": group["label"],
                "model_ids": [item["model_id"] for item in buckets],
                "token_type": preferred.get("token_type") or "",
                "remaining_fraction": preferred.get("remaining_fraction")
                if preferred.get("remaining_fraction") is not None
                else (min(fraction_values) if fraction_values else None),
                "remaining_amount": preferred.get("remaining_amount")
                if preferred.get("remaining_amount") is not None
                else (min(amount_values) if amount_values else None),
                "reset_time": preferred.get("reset_time") or (reset_values[0] if reset_values else ""),
            }
        )
    return grouped


def _gemini_cli_credit_balance(payload: Dict[str, Any]) -> Optional[float]:
    paid_tier = payload.get("paidTier") if isinstance(payload.get("paidTier"), dict) else payload.get("paid_tier")
    current_tier = payload.get("currentTier") if isinstance(payload.get("currentTier"), dict) else payload.get("current_tier")
    tier_info = paid_tier if isinstance(paid_tier, dict) else current_tier if isinstance(current_tier, dict) else {}
    credits = tier_info.get("availableCredits") if isinstance(tier_info.get("availableCredits"), list) else tier_info.get("available_credits")
    if not isinstance(credits, list):
        return None
    total = 0.0
    found = False
    for credit in credits:
        if not isinstance(credit, dict):
            continue
        credit_type = _value_as_text(credit.get("creditType") or credit.get("credit_type"))
        if credit_type != GEMINI_CLI_CREDIT_TYPE:
            continue
        amount = _value_as_number(credit.get("creditAmount") or credit.get("credit_amount"))
        if amount is None:
            continue
        total += amount
        found = True
    return total if found else None


def _gemini_cli_tier_info(payload: Dict[str, Any]) -> tuple[str, str]:
    paid_tier = payload.get("paidTier") if isinstance(payload.get("paidTier"), dict) else payload.get("paid_tier")
    current_tier = payload.get("currentTier") if isinstance(payload.get("currentTier"), dict) else payload.get("current_tier")
    tier_info = paid_tier if isinstance(paid_tier, dict) else current_tier if isinstance(current_tier, dict) else {}
    tier_id = _value_as_text(tier_info.get("id")).lower()
    tier_label = GEMINI_CLI_TIER_LABELS.get(tier_id, tier_id.replace("-", " ").title() if tier_id else "")
    return tier_id, tier_label


def _build_local_quota_payload() -> Dict[str, Any]:
    entries = _load_local_cli_proxy_auth_entries()
    if not entries:
        return {
            "available": False,
            "reason": "No local Gemini auth file was found in Research Companion.",
        }

    entry = entries[0]
    access_token = _resolve_local_cli_proxy_access_token(entry)
    if not access_token:
        return {
            "available": False,
            "reason": "Local Gemini auth token is missing or expired. Sign in again in ResearchCompanion.exe.",
            "email": entry.get("email", ""),
            "project_id": entry.get("project_id", ""),
        }

    project_id = _value_as_text(entry.get("project_id"))
    if not project_id:
        return {
            "available": False,
            "reason": "The local Gemini auth file does not include a project ID yet.",
            "email": entry.get("email", ""),
        }

    quota_payload = _local_cli_proxy_json_request(entry, GEMINI_CLI_QUOTA_URL, {"project": project_id})
    tier_payload = _local_cli_proxy_json_request(
        entry,
        GEMINI_CLI_LOAD_CODE_ASSIST_URL,
        {
            "cloudaicompanionProject": project_id,
            "metadata": {
                "ideType": "IDE_UNSPECIFIED",
                "platform": "PLATFORM_UNSPECIFIED",
                "pluginType": "GEMINI",
                "duetProject": project_id,
            },
        },
    )

    raw_buckets = quota_payload.get("buckets") if isinstance(quota_payload.get("buckets"), list) else []
    groups = _gemini_cli_bucket_groups(raw_buckets)
    tier_id, tier_label = _gemini_cli_tier_info(tier_payload)
    credit_balance = _gemini_cli_credit_balance(tier_payload)
    fractions = [group["remaining_fraction"] for group in groups if group.get("remaining_fraction") is not None]
    summary_fraction = min(fractions) if fractions else None

    return {
        "available": True,
        "email": entry.get("email", ""),
        "project_id": project_id,
        "auth_file": entry.get("name", ""),
        "auth_file_count": len(entries),
        "summary_fraction": summary_fraction,
        "tier_id": tier_id,
        "tier_label": tier_label,
        "premium_tier": tier_id in GEMINI_CLI_PREMIUM_TIER_IDS if tier_id else False,
        "credit_balance": credit_balance,
        "buckets": groups,
    }


def _cli_proxy_runtime_credentials() -> Any:
    entries = _load_local_cli_proxy_auth_entries()
    if not entries:
        return None
    return _cli_proxy_credentials_from_token_info(entries[0].get("token") if isinstance(entries[0].get("token"), dict) else {})


def _runtime_google_credentials() -> Any:
    if _local_llm_auth_mode() != "cli_proxy_oauth":
        return None
    return _cli_proxy_runtime_credentials()


def _summarize_image_bytes(
    *,
    image_bytes: bytes,
    mime_type: str,
    filename: str,
    credentials: Any,
) -> str:
    if credentials is None:
        return ""
    try:
        from google import genai
        from google.genai import types
    except Exception:
        return ""
    credentials, resolved_project, _source = resolve_vertex_runtime_credentials(credentials, project=GOOGLE_CLOUD_PROJECT)
    if credentials is None or not str(resolved_project or "").strip():
        return ""
    client = genai.Client(
        vertexai=True,
        credentials=credentials,
        project=resolved_project,
        location=GOOGLE_CLOUD_LOCATION or "global",
    )
    response = client.models.generate_content(
        model=DEFAULT_GEMINI_MODEL,
        contents=[
            types.Content(
                role="user",
                parts=[
                    types.Part.from_text(
                        text=(
                            "Read this uploaded image carefully and produce a concise research-oriented summary. "
                            "If it is a chart, table, figure, manuscript screenshot, or scanned page, extract the key facts, labels, "
                            "methods, and findings that would help an academic writing workflow. Keep it factual."
                        )
                    ),
                    types.Part.from_bytes(data=image_bytes, mime_type=mime_type or "image/png"),
                ],
            )
        ],
    )
    text = str(getattr(response, "text", "") or "").strip()
    if text:
        return _truncate_chars(text, 8000)
    return f"Uploaded image: {filename}"


def _attachment_context(records: List[Dict[str, Any]]) -> str:
    parts: List[str] = []
    for item in records:
        name = str(item.get("filename") or item.get("title") or "attachment").strip()
        file_type = str(item.get("file_type") or "").strip()
        extracted = str(item.get("extracted_text") or item.get("summary") or "").strip()
        if not extracted:
            continue
        parts.append(
            f"Attachment: {name}\nType: {file_type}\nContent excerpt:\n{_truncate_chars(extracted, 9000)}"
        )
    return "\n\n---\n\n".join(parts)


def _prompt_requests_revision(prompt: str) -> bool:
    text = str(prompt or "").lower()
    tokens = [
        "edit",
        "revise",
        "rewrite",
        "improve",
        "polish",
        "modify",
        "correct",
        "fix",
        "expand",
        "shorten",
        "based on attached",
        "dựa trên file",
        "chỉnh sửa",
        "viết lại",
        "sửa bài",
    ]
    return any(token in text for token in tokens)


def _select_source_manuscript(prompt: str, records: List[Dict[str, Any]]) -> str:
    if not _prompt_requests_revision(prompt):
        return ""
    candidates = []
    for item in records:
        text = str(item.get("extracted_text") or item.get("summary") or "").strip()
        if len(text) < 1200:
            continue
        filename = str(item.get("filename") or "").lower()
        score = len(text)
        if any(token in filename for token in ["draft", "paper", "article", "manuscript", "bai_bao", "nghien_cuu"]):
            score += 5000
        candidates.append((score, text))
    candidates.sort(key=lambda pair: pair[0], reverse=True)
    return _truncate_chars(candidates[0][1], 32000) if candidates else ""


_RESEARCH_QUERY_STOPWORDS = {
    "a", "an", "and", "are", "article", "attached", "attachment", "attachments", "based", "bai", "bao",
    "binh", "comment", "comments", "content", "cu", "cua", "draft", "edit", "editor", "execution", "expand",
    "file", "files", "fix", "follow", "for", "from", "improve", "into", "its", "log", "manuscript", "modify",
    "node", "paper", "phản", "phan", "progress", "queued", "reader", "refine", "researcher", "response",
    "review", "reviewer", "sua", "summary", "task", "that", "the", "them", "they", "this", "thu", "workflow",
    "writer", "your", "you", "dựa", "dua", "trên", "tren", "chỉnh", "chinh", "sửa", "viet", "viết", "lại",
    "lai", "bình", "luận", "luan", "phản", "biện", "bien", "cua", "bản", "thảo", "thao", "đủ", "du",
    "thông", "tin", "tác", "giả", "tiep", "tục", "continue",
}
_REVISION_ATTACHMENT_IGNORES = (
    "review",
    "comment",
    "referee",
    "feedback",
    "reply",
    "response",
    "bình_luận",
    "binh_luan",
    "phan_bien",
    "phản_biện",
)


def _clean_research_text(text: str) -> str:
    value = str(text or "").strip()
    if not value:
        return ""
    value = re.sub(r"https?://\S+", " ", value)
    value = re.sub(r"\bworkflow (?:queued|completed|stopped|failed)\b.*", " ", value, flags=re.IGNORECASE)
    value = re.sub(r"\b(?:follow the execution log|node-by-node progress|they will be treated as offline context)\b.*", " ", value, flags=re.IGNORECASE)
    value = re.sub(r"\battached \d+ file\(s\)\b.*", " ", value, flags=re.IGNORECASE)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _transliterate_ascii(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(text or ""))
    return normalized.encode("ascii", "ignore").decode("ascii")


def _tokenize_research_terms(text: str) -> List[str]:
    lowered = _clean_research_text(text).lower()
    tokens = re.findall(r"[^\W_]{3,}", lowered, flags=re.UNICODE)
    out: List[str] = []
    seen: set[str] = set()
    for token in tokens:
        if token in _RESEARCH_QUERY_STOPWORDS or token.isdigit():
            continue
        if token in seen:
            continue
        seen.add(token)
        out.append(token)
    return out


def _extract_candidate_title(text: str) -> str:
    for raw_line in str(text or "").splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip(" -:\t")
        if not line or len(line) < 18 or len(line) > 220:
            continue
        if sum(char.isalpha() for char in line) < 10:
            continue
        if line.lower().startswith(("abstract", "tóm tắt", "keywords", "introduction", "mở đầu")):
            continue
        return line
    return ""


def _attachment_query_candidates(records: List[Dict[str, Any]]) -> List[str]:
    candidates: List[str] = []
    for item in records:
        filename = str(item.get("filename") or item.get("title") or "").strip()
        lowered_filename = filename.lower()
        extracted = str(item.get("extracted_text") or item.get("summary") or "").strip()
        title = _extract_candidate_title(extracted)
        if title:
            candidates.append(title)
        base_name = os.path.splitext(filename)[0]
        if base_name and not any(token in lowered_filename for token in _REVISION_ATTACHMENT_IGNORES):
            candidates.append(base_name.replace("_", " "))
    return candidates


def _build_research_query(prompt: str, records: List[Dict[str, Any]], chat_history: List[ChatMessage]) -> str:
    prompt_text = _clean_research_text(prompt)
    user_messages = [
        _clean_research_text(item.content)
        for item in chat_history[-8:]
        if str(item.role or "user").lower() == "user" and _clean_research_text(item.content)
    ]
    attachment_candidates = _attachment_query_candidates(records)
    source_manuscript = _select_source_manuscript(prompt_text, records)
    manuscript_title = _extract_candidate_title(source_manuscript)

    preferred_parts: List[str] = []
    if manuscript_title:
        preferred_parts.append(manuscript_title)
    preferred_parts.extend(attachment_candidates[:4])
    preferred_parts.extend(user_messages[-3:])
    if prompt_text:
        preferred_parts.append(prompt_text)

    combined = " ".join(part for part in preferred_parts if part).strip()
    if not combined:
        return ""

    if manuscript_title and len(manuscript_title) <= 180:
        return manuscript_title

    tokens = _tokenize_research_terms(combined)
    if not tokens:
        fallback = re.sub(r"[()\\[\\],;:]+", " ", combined)
        fallback = re.sub(r"\s+", " ", fallback).strip()
        return _truncate_chars(fallback, 180)

    ascii_tokens = _tokenize_research_terms(_transliterate_ascii(combined))
    merged_tokens: List[str] = []
    seen_tokens: set[str] = set()
    for token in tokens + ascii_tokens:
        lowered = token.lower()
        if lowered in seen_tokens:
            continue
        seen_tokens.add(lowered)
        merged_tokens.append(token)
    return _truncate_chars(" ".join(merged_tokens[:12]), 180)


def _effective_google_oauth_client_config_values() -> tuple[str, str, str, str]:
    if GOOGLE_OAUTH_CLIENT_ID and GOOGLE_OAUTH_CLIENT_SECRET:
        return (
            GOOGLE_OAUTH_CLIENT_ID,
            GOOGLE_OAUTH_CLIENT_SECRET,
            GOOGLE_OAUTH_PROJECT_ID,
            "env",
        )
    return ("", "", GOOGLE_OAUTH_PROJECT_ID, "missing")


def _oauth_store_uses_firestore() -> bool:
    return OAUTH_STORE_BACKEND == "firestore"


def _firestore_collection_name(kind: str) -> str:
    prefix = OAUTH_FIRESTORE_COLLECTION_PREFIX or "research_oauth"
    return f"{prefix}_{kind}"


def _firestore_client() -> Any:
    global _firestore_client_cache
    if not _oauth_store_uses_firestore():
        return None
    with _firestore_client_lock:
        if _firestore_client_cache is not None:
            return _firestore_client_cache
        from google.cloud import firestore

        _firestore_client_cache = firestore.Client()
        return _firestore_client_cache


def _ensure_oauth_store_ready() -> None:
    if not _oauth_store_uses_firestore():
        return
    try:
        _firestore_client()
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Firestore OAuth store is not ready: {exc}",
        ) from exc


def _firestore_get_document(kind: str, doc_id: str) -> Optional[Dict[str, Any]]:
    if not doc_id:
        return None
    client = _firestore_client()
    if client is None:
        return None
    snapshot = client.collection(_firestore_collection_name(kind)).document(doc_id).get()
    if not snapshot.exists:
        return None
    data = snapshot.to_dict() or {}
    return data if isinstance(data, dict) else None


def _firestore_find_one(kind: str, field_name: str, value: Any) -> Optional[Dict[str, Any]]:
    client = _firestore_client()
    if client is None:
        return None
    from google.cloud.firestore_v1.base_query import FieldFilter

    query = (
        client.collection(_firestore_collection_name(kind))
        .where(filter=FieldFilter(field_name, "==", value))
        .limit(1)
    )
    for snapshot in query.stream():
        data = snapshot.to_dict() or {}
        return data if isinstance(data, dict) else None
    return None


def _firestore_set_document(kind: str, doc_id: str, data: Dict[str, Any]) -> None:
    client = _firestore_client()
    if client is None:
        return
    client.collection(_firestore_collection_name(kind)).document(doc_id).set(data)


def _firestore_delete_document(kind: str, doc_id: str) -> None:
    client = _firestore_client()
    if client is None or not doc_id:
        return
    client.collection(_firestore_collection_name(kind)).document(doc_id).delete()


def _write_auth_file_token_info(auth_file: str, email: str, token_info: Dict[str, Any]) -> None:
    if not auth_file:
        return
    try:
        write_json_file(
            auth_file,
            {
                "provider": "gemini",
                "auth_type": "google_oauth",
                "email": email,
                "token_info": token_info,
            },
        )
    except Exception:
        return


def _google_credentials_from_token_info(
    token_info: Any,
    *,
    on_refresh: Optional[Any] = None,
) -> Any:
    if not isinstance(token_info, dict):
        return None
    try:
        from google.auth.transport.requests import Request as GoogleAuthRequest
        from google.oauth2.credentials import Credentials

        credentials = Credentials.from_authorized_user_info(token_info, scopes=GOOGLE_GEMINI_OAUTH_SCOPES)
        if credentials.expired and credentials.refresh_token:
            credentials.refresh(GoogleAuthRequest())
            refreshed = json.loads(credentials.to_json())
            if callable(on_refresh):
                on_refresh(refreshed)
        return credentials
    except Exception:
        return None


def _session_credentials_from_entry(session_id: str, entry: Dict[str, Any]) -> Any:
    email = str(entry.get("email") or "").strip()
    auth_file = str(entry.get("auth_file") or "").strip()
    token_info = entry.get("token_info") if isinstance(entry.get("token_info"), dict) else None

    if token_info:
        def _persist(updated_token_info: Dict[str, Any]) -> None:
            refreshed_entry = dict(entry)
            refreshed_entry["token_info"] = updated_token_info
            refreshed_entry["updated_at"] = time.time()
            _upsert_google_session(session_id, refreshed_entry)
            _write_auth_file_token_info(auth_file, email, updated_token_info)

        credentials = _google_credentials_from_token_info(token_info, on_refresh=_persist)
        if credentials is not None:
            return credentials

    if auth_file:
        credentials = google_credentials_from_auth_file(auth_file)
        if credentials is not None:
            try:
                updated = json.loads(credentials.to_json())
                refreshed_entry = dict(entry)
                refreshed_entry["token_info"] = updated
                refreshed_entry["updated_at"] = time.time()
                _upsert_google_session(session_id, refreshed_entry)
            except Exception:
                pass
            return credentials

    return None


def _google_client_config() -> Dict[str, Any]:
    client_id, client_secret, project_id, _source = _effective_google_oauth_client_config_values()
    if client_id and client_secret:
        return {
            "web": {
                "client_id": client_id,
                "client_secret": client_secret,
                "project_id": project_id,
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
                "redirect_uris": [f"{PUBLIC_BACKEND_URL}/oauth/google/callback"],
            }
        }

    raise RuntimeError(
        "Google OAuth client is not configured. Set GOOGLE_OAUTH_CLIENT_ID and GOOGLE_OAUTH_CLIENT_SECRET on the backend."
    )


def _google_client_section() -> Dict[str, Any]:
    config = _google_client_config()
    section = config.get("web") or config.get("installed") or {}
    if not isinstance(section, dict):
        raise RuntimeError("Google OAuth client config is invalid.")
    return section


def _google_redirect_uri() -> str:
    section = _google_client_section()
    redirect_uris = section.get("redirect_uris") or []
    if isinstance(redirect_uris, list) and redirect_uris:
        return str(redirect_uris[0] or "").strip()
    return f"{PUBLIC_BACKEND_URL}/oauth/google/callback"


def _oauth_success_html(message: str, *, success: bool) -> str:
    color = "#16a34a" if success else "#dc2626"
    title = "Google OAuth Completed" if success else "Google OAuth Failed"
    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{title}</title>
  <style>
    body {{ font-family: Arial, sans-serif; background: #0f172a; color: #e5e7eb; padding: 2rem; }}
    .box {{ max-width: 680px; margin: 2rem auto; padding: 1.25rem 1.4rem; border: 1px solid #334155; border-radius: 14px; background: #111827; }}
    .status {{ color: {color}; font-weight: 700; margin-bottom: 0.75rem; }}
  </style>
</head>
<body>
  <div class="box">
    <div class="status">{html.escape(title)}</div>
    <div>{html.escape(message)}</div>
    <p>You can close this tab and return to the app.</p>
  </div>
  <script>
    try {{
      if (window.opener) {{
        window.opener.postMessage({{"type": "google-oauth-finished"}}, "{html.escape(FRONTEND_ORIGIN)}");
      }}
    }} catch (_err) {{}}
    setTimeout(function () {{
      try {{ window.close(); }} catch (_err) {{}}
    }}, 1200);
  </script>
</body>
</html>"""


def _oauth_google_credentials_from_request(
    request: Optional[Request],
    *,
    oauth_session_id: str = "",
) -> Any:
    session_id = str(oauth_session_id or "").strip()
    if request is not None and not session_id:
        session_id = str(request.headers.get("x-oauth-session") or "").strip()
    if not session_id:
        return None

    entry = _google_session_entry(session_id)
    if not entry:
        return None
    return _session_credentials_from_entry(session_id, entry)


def _google_session_entry(session_id: str) -> Optional[Dict[str, Any]]:
    if not session_id:
        return None
    if _oauth_store_uses_firestore():
        return _firestore_get_document("google_sessions", session_id)
    sessions = _read_web_oauth_sessions().get("sessions") or {}
    entry = sessions.get(session_id) if isinstance(sessions, dict) else None
    return entry if isinstance(entry, dict) else None


def _upsert_google_session(session_id: str, entry: Dict[str, Any]) -> None:
    if _oauth_store_uses_firestore():
        payload = dict(entry)
        payload["oauth_session_id"] = session_id
        _firestore_set_document("google_sessions", session_id, payload)
        return
    store = _read_web_oauth_sessions()
    sessions = store.setdefault("sessions", {})
    if not isinstance(sessions, dict):
        sessions = {}
        store["sessions"] = sessions
    sessions[session_id] = entry
    _write_web_oauth_sessions(store)


def _delete_google_session(session_id: str) -> None:
    if not session_id:
        return
    if _oauth_store_uses_firestore():
        _firestore_delete_document("google_sessions", session_id)
        return
    store = _read_web_oauth_sessions()
    sessions = store.get("sessions") or {}
    if not isinstance(sessions, dict) or session_id not in sessions:
        return
    del sessions[session_id]
    store["sessions"] = sessions
    _write_web_oauth_sessions(store)


def _google_auth_url_for_flow(_flow_id: str, state_token: str) -> tuple[str, str]:
    client_section = _google_client_section()
    client_id = str(client_section.get("client_id") or "").strip()
    client_secret = str(client_section.get("client_secret") or "").strip()
    auth_uri = str(client_section.get("auth_uri") or "https://accounts.google.com/o/oauth2/v2/auth").strip()
    if not client_id:
        raise RuntimeError("Google OAuth client_id is missing.")
    if not client_secret:
        raise RuntimeError("Google OAuth client_secret is missing.")

    params = {
        "access_type": "offline",
        "client_id": client_id,
        "prompt": "consent",
        "redirect_uri": _google_redirect_uri(),
        "response_type": "code",
        "scope": " ".join(GOOGLE_GEMINI_OAUTH_SCOPES),
        "state": state_token,
    }
    return auth_uri + "?" + urllib.parse.urlencode(params), client_id


def _oauth_flow(flow_id: str = "", state_token: str = "") -> Optional[Dict[str, Any]]:
    if _oauth_store_uses_firestore():
        if flow_id:
            return _firestore_get_document("google_flows", flow_id)
        if state_token:
            return _firestore_find_one("google_flows", "state", state_token)
        return None
    flows = _read_web_oauth_state().get("flows") or {}
    if not isinstance(flows, dict):
        return None
    if flow_id:
        entry = flows.get(flow_id)
        return entry if isinstance(entry, dict) else None
    for entry in flows.values():
        if isinstance(entry, dict) and str(entry.get("state") or "") == state_token:
            return entry
    return None


def _upsert_oauth_flow(flow_id: str, entry: Dict[str, Any]) -> None:
    if _oauth_store_uses_firestore():
        payload = dict(entry)
        payload["flow_id"] = flow_id
        _firestore_set_document("google_flows", flow_id, payload)
        return
    store = _read_web_oauth_state()
    flows = store.setdefault("flows", {})
    if not isinstance(flows, dict):
        flows = {}
        store["flows"] = flows
    flows[flow_id] = entry
    _write_web_oauth_state(store)


def _update_oauth_flow(flow_id: str, **changes: Any) -> Optional[Dict[str, Any]]:
    if _oauth_store_uses_firestore():
        entry = _firestore_get_document("google_flows", flow_id)
        if not isinstance(entry, dict):
            return None
        entry.update(changes)
        entry["updated_at"] = time.time()
        entry["flow_id"] = flow_id
        _firestore_set_document("google_flows", flow_id, entry)
        return entry
    store = _read_web_oauth_state()
    flows = store.get("flows") or {}
    entry = flows.get(flow_id) if isinstance(flows, dict) else None
    if not isinstance(entry, dict):
        return None
    entry.update(changes)
    entry["updated_at"] = time.time()
    flows[flow_id] = entry
    store["flows"] = flows
    _write_web_oauth_state(store)
    return entry


def _job_to_dict(job: JobSnapshot) -> Dict[str, Any]:
    return job.model_dump()


def _truncate_text(value: Any, limit: int) -> str:
    text = str(value or "")
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 1)].rstrip() + "…"


def _sanitize_job_payload(data: Dict[str, Any]) -> Dict[str, Any]:
    payload = dict(data)

    request_payload = payload.get("request") if isinstance(payload.get("request"), dict) else {}
    if isinstance(request_payload, dict):
        chat_history = request_payload.get("chat_history") if isinstance(request_payload.get("chat_history"), list) else []
        payload["request"] = {
            "topic": _truncate_text(request_payload.get("topic", ""), 3000),
            "language": request_payload.get("language", ""),
            "provider": request_payload.get("provider", ""),
            "model": _truncate_text(request_payload.get("model", ""), 80),
            "target_word_count": request_payload.get("target_word_count", 0),
            "reference_target": request_payload.get("reference_target", 0),
            "attachment_ids": [str(item)[:80] for item in (request_payload.get("attachment_ids") or [])[:20]],
            "chat_history": [
                {
                    "role": _truncate_text(entry.get("role", ""), 20),
                    "content": _truncate_text(entry.get("content", ""), 600),
                }
                for entry in chat_history[-10:]
                if isinstance(entry, dict)
            ],
        }

    papers: List[Dict[str, Any]] = []
    for item in payload.get("papers", []) or []:
        if not isinstance(item, dict):
            continue
        papers.append(
            {
                "title": _truncate_text(item.get("title", ""), 600),
                "authors": _truncate_text(item.get("authors", ""), 300),
                "year": item.get("year", 0),
                "doi": _truncate_text(item.get("doi", ""), 120),
                "abstract": _truncate_text(item.get("abstract", ""), 1600),
                "source_db": _truncate_text(item.get("source_db", ""), 120),
                "landing_url": _truncate_text(item.get("landing_url", ""), 500),
                "access_date": _truncate_text(item.get("access_date", ""), 64),
                "search_rank": item.get("search_rank", 0),
                "relevance_score": item.get("relevance_score", 0),
            }
        )
    payload["papers"] = papers[:80]

    attachments: List[Dict[str, Any]] = []
    for item in payload.get("attachments", []) or []:
        if not isinstance(item, dict):
            continue
        attachments.append(
            {
                "id": _truncate_text(item.get("id", ""), 80),
                "filename": _truncate_text(item.get("filename", ""), 200),
                "file_type": _truncate_text(item.get("file_type", ""), 40),
                "preview": _truncate_text(item.get("preview", ""), 400),
            }
        )
    payload["attachments"] = attachments[:40]

    payload.pop("texts", None)
    payload.pop("quotations", None)

    payload["logs"] = [_truncate_text(line, 1200) for line in (payload.get("logs", []) or [])][-500:]
    payload["manager_output"] = _truncate_text(payload.get("manager_output", ""), 30000)
    payload["outline"] = _truncate_text(payload.get("outline", ""), 30000)
    payload["review_feedback"] = _truncate_text(payload.get("review_feedback", ""), 50000)
    payload["qa_summary"] = _truncate_text(payload.get("qa_summary", ""), 30000)
    payload["draft"] = _truncate_text(payload.get("draft", ""), 180000)
    payload["final_markdown"] = _truncate_text(payload.get("final_markdown", ""), 180000)

    return payload


def _jobs_store_uses_firestore() -> bool:
    return _oauth_store_uses_firestore()


def _job_from_dict(data: Dict[str, Any]) -> Optional[JobSnapshot]:
    if not isinstance(data, dict):
        return None
    try:
        return JobSnapshot.model_validate(_sanitize_job_payload(data))
    except Exception:
        return None


def _job_entry(job_id: str) -> Optional[JobSnapshot]:
    if not job_id:
        return None
    if _jobs_store_uses_firestore():
        loaded = _job_from_dict(_firestore_get_document("jobs", job_id) or {})
        if loaded is not None:
            with _jobs_lock:
                _jobs[job_id] = loaded
            return loaded
    with _jobs_lock:
        job = _jobs.get(job_id)
        return job


def _persist_job(job: JobSnapshot) -> None:
    if not _jobs_store_uses_firestore():
        return
    _firestore_set_document("jobs", job.id, _sanitize_job_payload(_job_to_dict(job)))


def _delete_job_entry(job_id: str) -> None:
    if _jobs_store_uses_firestore():
        _firestore_delete_document("jobs", job_id)


def _update_job(job_id: str, **changes: Any) -> None:
    with _jobs_lock:
        job = _jobs.get(job_id)
        if job is None and _jobs_store_uses_firestore():
            loaded = _job_from_dict(_firestore_get_document("jobs", job_id) or {})
            if loaded is not None:
                _jobs[job_id] = loaded
                job = loaded
        if job is None:
            return
        for key, value in changes.items():
            setattr(job, key, value)
        job.updated_at = time.time()
        snapshot = JobSnapshot.model_validate(job.model_dump())
    _persist_job(snapshot)


def _append_log(job_id: str, line: str) -> None:
    with _jobs_lock:
        job = _jobs.get(job_id)
        if job is None and _jobs_store_uses_firestore():
            loaded = _job_from_dict(_firestore_get_document("jobs", job_id) or {})
            if loaded is not None:
                _jobs[job_id] = loaded
                job = loaded
        if job is None:
            return
        job.logs.append(line)
        job.updated_at = time.time()
        snapshot = JobSnapshot.model_validate(job.model_dump())
    _persist_job(snapshot)


def _job_stop_event(job_id: str) -> threading.Event:
    with _job_stop_lock:
        event = _job_stop_events.get(job_id)
        if event is None:
            event = threading.Event()
            _job_stop_events[job_id] = event
        return event


def _job_is_stopping(job_id: str) -> bool:
    if _job_stop_event(job_id).is_set():
        return True
    if _jobs_store_uses_firestore():
        entry = _firestore_get_document("jobs", job_id)
        if isinstance(entry, dict) and str(entry.get("status") or "").strip().lower() == "cancelling":
            return True
    return False


def _clear_job_stop(job_id: str) -> None:
    with _job_stop_lock:
        _job_stop_events.pop(job_id, None)


def _raise_if_stopped(job_id: str) -> None:
    if _job_is_stopping(job_id):
        raise JobCancelledError("Workflow stopped by user.")


def _is_core_editor_log(line: str) -> bool:
    normalized = (line or "").lower()
    return "editor started:" in normalized or "editor done:" in normalized


def _sanitize_online_statuses(statuses: Dict[str, str]) -> Dict[str, str]:
    sanitized = dict(statuses)
    sanitized["Reviewer"] = sanitized.get("Reviewer", "Pending")
    sanitized["Editor"] = "Pending"
    sanitized["Translator"] = "Pending"
    return sanitized


def _make_hooks(job_id: str, *, suppress_core_editor: bool = False) -> WorkflowHooks:
    def on_status(statuses: Dict[str, str]) -> None:
        _raise_if_stopped(job_id)
        safe_statuses = _sanitize_online_statuses(statuses) if suppress_core_editor else dict(statuses)
        _update_job(job_id, node_statuses=safe_statuses)

    def on_log(_logs: List[str], line: str) -> None:
        _raise_if_stopped(job_id)
        if suppress_core_editor and _is_core_editor_log(line):
            return
        _append_log(job_id, line)

    def on_manager(text: str) -> None:
        _raise_if_stopped(job_id)
        _update_job(job_id, manager_output=text)

    def on_outline(text: str) -> None:
        _raise_if_stopped(job_id)
        _update_job(job_id, outline=text)

    def on_metadata(items: List[Dict[str, Any]]) -> None:
        _raise_if_stopped(job_id)
        _update_job(job_id, papers=items)

    def on_reader_data(texts: List[Dict[str, Any]], quotes: List[Dict[str, Any]]) -> None:
        _raise_if_stopped(job_id)
        return

    def on_draft(text: str) -> None:
        _raise_if_stopped(job_id)
        _update_job(job_id, draft=text)

    def on_formatted(text: str) -> None:
        _raise_if_stopped(job_id)
        if suppress_core_editor:
            return
        _update_job(job_id, final_markdown=text)

    return WorkflowHooks(
        on_status=on_status,
        on_log=on_log,
        on_manager=on_manager,
        on_outline=on_outline,
        on_metadata=on_metadata,
        on_reader_data=on_reader_data,
        on_draft=on_draft,
        on_formatted=on_formatted,
        should_stop=lambda: _raise_if_stopped(job_id),
    )


def _strip_references_section(text: str) -> str:
    if not text:
        return ""
    parts = re.split(r"\n##\s+References\b", text, maxsplit=1, flags=re.IGNORECASE)
    return parts[0].strip() if parts else text.strip()


def _online_reviewer_prompt(
    *,
    language: str,
    draft: str,
    target_word_count: int,
    reference_target: int,
) -> str:
    current_words = count_words(draft)
    min_words = max(300, int(target_word_count * 0.9))
    max_words = max(min_words + 80, int(target_word_count * 1.08))
    return f"""Revise the academic manuscript body below in {language}.

Current total body length: about {current_words} words.
Target total body length after revision: {min_words}-{max_words} words, aiming for about {target_word_count} words.
Maximum number of distinct cited sources allowed in the body: {reference_target}.

Requirements:
- Preserve IMRaD structure and academic tone.
- If the draft is too long, compress it decisively.
- If the draft is too short, expand modestly but only using claims already supported by citations already present.
- Do not invent any new citations or sources.
- Reduce citation spread so the body relies on no more than {reference_target} distinct cited sources.
- Return Markdown body only.
- Do not include a References section.

BODY:
{draft}
"""


def _online_reviewer_feedback_prompt(
    *,
    language: str,
    draft: str,
    target_word_count: int,
    reference_target: int,
    manager_brief: str = "",
) -> str:
    current_words = count_words(draft)
    return f"""Review the academic draft below in {language}.

Current body length: about {current_words} words.
Target body length: about {target_word_count} words.
Maximum distinct cited sources allowed in the body: {reference_target}.

Return Markdown only with these sections:
## Strengths
## Problems
## Required Revisions

Focus on:
- logical coherence and argument quality
- weak sections or repetition
- whether claims are overextended relative to citations
- whether the number of cited sources should be reduced
- whether the body is too long or too short for the target
- Management guidance: {manager_brief or "(none)"}

DRAFT:
{draft}
"""


def _online_rewrite_prompt(
    *,
    language: str,
    draft: str,
    review_feedback: str,
    target_word_count: int,
    reference_target: int,
    manager_brief: str = "",
) -> str:
    current_words = count_words(draft)
    min_words = max(280, int(target_word_count * 0.9))
    max_words = max(min_words + 80, int(target_word_count * 1.08))
    return f"""Rewrite the academic manuscript body below in {language} by following the REVIEW FEEDBACK.

Current body length: about {current_words} words.
Required body length after rewrite: {min_words}-{max_words} words, aiming for about {target_word_count} words.
Maximum number of distinct cited sources allowed in the body: {reference_target}.

Requirements:
- Preserve IMRaD structure.
- Apply the review feedback seriously.
- Strengthen weak reasoning and remove repetition.
- Do not invent new citations or sources.
- Reduce citation spread to at most {reference_target} distinct cited sources.
- Return Markdown body only.
- Do not include a References section.
- Management guidance: {manager_brief or "(none)"}

REVIEW FEEDBACK:
{review_feedback}

DRAFT:
{draft}
"""


def _online_abstract_prompt(
    *,
    language: str,
    manuscript_body: str,
    target_word_count: int,
    manager_brief: str = "",
    source_language: str = "",
) -> str:
    abstract_target = min(250, max(120, int(target_word_count * 0.12)))
    translation_instruction = (
        f"Write the abstract in {language} even if the manuscript body is currently in {source_language}. "
        "Translate faithfully while preserving the academic meaning and all citation markers."
        if source_language and source_language != language
        else ""
    )
    return f"""Write an academic abstract in {language} for the manuscript body below.

Target abstract length: about {abstract_target} words.

Requirements:
- Output Markdown only.
- Start with `## Abstract`.
- Summarize background, objective, approach, key findings, and implications.
- Do not include a References section.
- Do not invent evidence beyond the manuscript body.
- Avoid bullet points.
- {translation_instruction or f"Keep the abstract in {language}."}
- Management guidance: {manager_brief or "(none)"}

MANUSCRIPT BODY:
{manuscript_body}
"""


def _online_editor_prompt(
    *,
    language: str,
    draft: str,
    target_word_count: int,
    manager_brief: str = "",
    source_language: str = "",
) -> str:
    min_words = max(300, int(target_word_count * 0.9))
    max_words = max(min_words + 80, int(target_word_count * 1.08))
    translation_instruction = (
        f"Translate the manuscript from {source_language} into {language} while polishing it."
        if source_language and source_language != language
        else f"Keep the manuscript in {language}."
    )
    return f"""Polish the academic manuscript body below in {language}.

Requirements:
- Keep the IMRaD structure.
- Preserve all existing in-text citations.
- Keep the final body length within {min_words}-{max_words} words.
- Improve clarity, flow, and coherence.
- {translation_instruction}
- Return Markdown body only.
- Do not include a References section.
- Management guidance: {manager_brief or "(none)"}

BODY:
{draft}
"""


def _online_translator_prompt(
    *,
    language: str,
    draft: str,
    source_language: str,
    manager_brief: str = "",
    attachment_context: str = "",
) -> str:
    glossary_context = attachment_context[:12000].strip()
    return f"""Translate the academic manuscript below from {source_language} into {language}.

Requirements:
- Preserve Markdown headings, structure, in-text citation markers, years, numbers, and factual meaning exactly.
- Use attached glossary, spreadsheet, manuscript, and offline context whenever they define preferred terminology or phrasing.
- Prefer academically natural terminology in {language}; keep only unavoidable proper nouns or official titles in their original language.
- Do not add a References section.
- Return Markdown body only.
- Management guidance: {manager_brief or "(none)"}

Offline terminology context:
{glossary_context or "(none)"}

MANUSCRIPT:
{draft}
"""


def _final_translation_prompt(*, source_language: str, target_language: str, draft: str) -> str:
    return f"""Translate the academic manuscript below from {source_language} to {target_language}.

Requirements:
- Preserve Markdown headings and overall IMRaD structure.
- Preserve all in-text citation markers and years exactly.
- Keep the manuscript academically natural in {target_language}.
- All prose paragraphs must be in {target_language}. Do not leave explanatory sentences in {source_language}.
- Keep only proper nouns, dataset names, institutional names, or official document titles in their original language when necessary.
- Do not add a References section.
- Return Markdown body only.

MANUSCRIPT:
{draft}
"""


def _heading_lines(markdown_text: str) -> List[str]:
    return [line.strip() for line in str(markdown_text or "").splitlines() if re.match(r"^\s{0,3}#{2,6}\s+\S+", line)]


def _build_quality_summary(
    *,
    final_markdown: str,
    effective_target_word_count: int,
    effective_reference_target: int,
    output_language: str,
    task_mode: str,
) -> str:
    body = _strip_references_section(final_markdown)
    headings = _heading_lines(body)
    reference_lines = [line for line in str(final_markdown or "").splitlines() if line.strip().startswith("- ")]
    body_words = count_words(body)
    detected_language = guess_language_label(body[:5000])
    within_range = max(300, int(effective_target_word_count * 0.88)) <= body_words <= max(380, int(effective_target_word_count * 1.12))

    summary_lines = [
        "## Quality Check",
        "",
        f"- Task mode: `{task_mode or 'unknown'}`",
        f"- Output language requested: `{output_language}`",
        f"- Output language detected: `{detected_language}`",
        f"- Body word count: `{body_words}`",
        f"- Target word count: `{effective_target_word_count}`",
        f"- Word-count check: `{'pass' if within_range else 'review needed'}`",
        f"- References in final section: `{len(reference_lines)}`",
        f"- Reference target: `{effective_reference_target}`",
        "",
        "## Structure Check",
        "",
    ]
    if headings:
        summary_lines.extend([f"- {line}" for line in headings[:18]])
    else:
        summary_lines.append("- No Markdown section headings were detected.")
    return "\n".join(summary_lines).strip()


def _markdown_to_docx_bytes(markdown_text: str) -> bytes:
    from docx import Document

    document = Document()
    lines = str(markdown_text or "").splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            text = heading_match.group(2).strip()
            document.add_heading(text, level=level if level <= 4 else 4)
            continue
        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue
        document.add_paragraph(stripped)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _safe_docx_filename(value: str) -> str:
    text = re.sub(r"[^a-zA-Z0-9._-]+", "-", str(value or "").strip()).strip("._-")
    if not text:
        text = "research-workflow"
    if not text.lower().endswith(".docx"):
        text += ".docx"
    return text


def _ensure_output_language(
    *,
    cfg: LLMConfig,
    draft: str,
    target_language: str,
    should_stop: Any,
    max_attempts: int = 2,
) -> str:
    current = str(draft or "").strip()
    if not current:
        return current

    detected = guess_language_label(current[:5000])
    if detected == target_language:
        return current

    for _attempt in range(max_attempts):
        current = (
            llm_complete_text(
                cfg=cfg,
                prompt=_final_translation_prompt(
                    source_language=detected,
                    target_language=target_language,
                    draft=current,
                ),
                should_stop=should_stop,
            ).strip()
            or current
        )
        detected = guess_language_label(current[:5000])
        if detected == target_language:
            break
    return current


def _filter_online_logs(logs: List[str]) -> List[str]:
    return [line for line in logs if not _is_core_editor_log(line)]


def _local_llm_auth_mode() -> str:
    if _is_local_companion_runtime() and _load_local_cli_proxy_auth_entries():
        auth_state = _read_local_cli_proxy_auth_state()
        if str(auth_state.get("status") or "").strip().lower() == "invalid":
            return "cli_proxy_reauth_required"
        return "cli_proxy_oauth"
    return "google_oauth"


def _resolve_models(
    google_credentials: Any = None,
) -> tuple[list[str], str, str]:
    auth_mode = _local_llm_auth_mode()
    models, source_note = get_provider_models(
        "Gemini",
        gemini_auth_mode=auth_mode,
        google_credentials=google_credentials,
    )
    default_model = default_model_for_provider(
        "Gemini",
        gemini_auth_mode=auth_mode,
        google_credentials=google_credentials,
    )
    return models, default_model, source_note


def _require_local_companion_oauth() -> None:
    if _local_llm_auth_mode() != "cli_proxy_oauth":
        raise HTTPException(
            status_code=400,
            detail="Local companion OAuth is required. Start ResearchCompanion.exe and sign in first.",
        )


def _resolve_attachment_records(attachment_ids: List[str]) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []
    for attachment_id in attachment_ids:
        record = _upload_record(str(attachment_id or "").strip())
        if isinstance(record, dict):
            records.append(record)
    return records


def _chat_messages_block(messages: List[ChatMessage], *, limit: int = 14) -> str:
    lines: List[str] = []
    for item in messages[-limit:]:
        role = str(item.role or "user").strip().title()
        content = str(item.content or "").strip()
        if content:
            lines.append(f"{role}: {content}")
    return "\n".join(lines)


def _extract_json_object(raw_text: str) -> Dict[str, Any]:
    text = str(raw_text or "").strip()
    if not text:
        return {}
    try:
        parsed = json.loads(text)
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not match:
            return {}
        try:
            parsed = json.loads(match.group(0))
            return parsed if isinstance(parsed, dict) else {}
        except Exception:
            return {}


def _workflow_chat_protocol_prompt(
    *,
    language: str,
    target_word_count: Optional[int],
    reference_target: Optional[int],
    messages: List[ChatMessage],
    attachment_context: str,
    source_manuscript: str,
    heuristic_research_query: str,
) -> str:
    target_word_count_label = str(target_word_count) if isinstance(target_word_count, int) else "Unknown - ask the user"
    reference_target_label = str(reference_target) if isinstance(reference_target, int) else "Unknown - ask the user"
    return f"""You are Gemini acting as the conversational front layer of an academic workflow system.

You must both:
1. reply naturally to the user as a chat assistant
2. decide whether the conversation is ready to trigger the writing workflow automatically

Return ONLY valid JSON with this exact schema:
{{
  "workflow_intent_detected": true,
  "task_mode": "new_paper | revise_manuscript | critique_and_rewrite | summarize | compare | literature_review | chat_only",
  "topic": "clean workflow topic or empty string",
  "target_word_count": 1234,
  "reference_target": 12,
  "missing_requirements": ["topic", "word_count", "reference_count"],
  "assistant_reply": "natural chat reply to the user",
  "should_start_workflow": true,
  "workflow_topic": "concise actionable workflow task",
  "reason": "short explanation",
  "readiness": "chat_only | needs_clarification | ready_for_workflow"
}}

Protocol:
- Read the entire conversation, not just the final user message.
- You are the workflow front controller. Infer the user's real task, required scope fields, and whether the workflow should run now.
- If the user is casually asking, brainstorming, or clarifying, set `should_start_workflow` to false.
- If the user clearly wants the system to perform academic work such as writing, revising, rewriting, responding to reviewer comments, analyzing attached files, comparing materials, or extracting evidence, and the required scope fields are already complete, set `should_start_workflow` to true immediately.
- Do NOT ask for an extra confirmation turn once the request is clear and the required scope fields are complete.
- If the user gives a direct execution signal such as `bắt đầu`, `chạy đi`, `xác nhận`, `go ahead`, `proceed`, or `start`, treat that as permission to run immediately.
- If information is still missing, use `readiness = "needs_clarification"` and ask only the smallest set of questions needed to proceed.
- You must fill `workflow_intent_detected`, `task_mode`, `topic`, `target_word_count`, `reference_target`, and `missing_requirements` from the conversation.
- For a new paper request, if `target_word_count` or `reference_target` is still unknown, put those fields in `missing_requirements` and ask for them first.
- For a new paper request, do NOT switch to deeper content questions until `target_word_count` and `reference_target` are both filled.
- If the user gives values like `1000 từ và 10 tài liệu`, `3000 words and 12 references`, or `15 citations`, map them to `target_word_count` and `reference_target`.
- For revision tasks based on an attached manuscript, keep `target_word_count` and `reference_target` null unless the user explicitly changes them.
- When the task is confirmed and ready, let `assistant_reply` sound like an orchestrator: briefly confirm that the source materials are sufficient, summarize the deliverable, and state that the workflow is being launched.
- `workflow_topic` must be a clean task instruction for the workflow manager. Do not include UI/log text or raw filenames unless needed for meaning.
- Keep `assistant_reply` concise and useful. If you trigger the workflow, say so explicitly.
- Allowed entries in `missing_requirements` are only: `topic`, `word_count`, `reference_count`.
- Critical launch rule: when a new paper request already has `topic`, `target_word_count`, and `reference_target` in the conversation, you must launch immediately and must not ask whether to start.
- Critical launch rule: if the latest user message provides the final missing field, set `should_start_workflow` to true in the same response.

Known scope extracted from the conversation so far:
- Language: {language}
- Target words: {target_word_count_label}
- Reference target: {reference_target_label}
- Heuristic research query: {heuristic_research_query or "(none)"}

Conversation:
{_chat_messages_block(messages)}

Attachment context:
{attachment_context[:9000] if attachment_context else "(none)"}

Source manuscript excerpt:
{source_manuscript[:7000] if source_manuscript else "(none)"}
"""


def _workflow_scope_extraction_prompt(
    *,
    messages: List[ChatMessage],
    source_manuscript: str,
    attachment_context: str,
    known_topic: str,
    known_target_word_count: Optional[int],
    known_reference_target: Optional[int],
) -> str:
    topic_label = known_topic or "Unknown"
    word_label = str(known_target_word_count) if isinstance(known_target_word_count, int) else "Unknown"
    ref_label = str(known_reference_target) if isinstance(known_reference_target, int) else "Unknown"
    return f"""You extract workflow requirements from a multi-turn academic writing conversation.

Return ONLY valid JSON with this schema:
{{
  "topic": "string or empty if unknown",
  "target_word_count": 1234,
  "reference_target": 12
}}

Rules:
- Read the entire conversation, not just the final user turn.
- Preserve any already-known values unless the conversation clearly overrides them.
- `topic` is the actual paper/revision topic, not a meta instruction.
- `target_word_count` must be an integer or null.
- `reference_target` must be an integer or null.
- If the user says things like `10 tài liệu`, `10 tài liệu tham khảo`, `10 references`, `10 citations`, or `10 nguồn`, map that to `reference_target = 10`.
- If a value is still missing or unclear, return null for that field.

Known values so far:
- Topic: {topic_label}
- Target words: {word_label}
- Reference target: {ref_label}

Conversation:
{_chat_messages_block(messages, limit=20)}

Attachment context:
{attachment_context[:4000] if attachment_context else "(none)"}

Source manuscript excerpt:
{source_manuscript[:3000] if source_manuscript else "(none)"}
"""


def _fold_text_for_match(text: str) -> str:
    lowered = normalize_for_match(str(text or "")).lower()
    folded = unicodedata.normalize("NFKD", lowered)
    folded = "".join(ch for ch in folded if unicodedata.category(ch) != "Mn")
    folded = folded.replace("đ", "d")
    return re.sub(r"\s+", " ", folded).strip()


def _assistant_requested_workflow_confirmation(messages: List[ChatMessage]) -> bool:
    last_assistant = ""
    for item in reversed(messages):
        if str(item.role or "").lower() == "assistant":
            last_assistant = str(item.content or "").strip()
            break
    if not last_assistant:
        return False
    folded = _fold_text_for_match(last_assistant)
    confirmation_markers = [
        "confirm",
        "confirmation",
        "should i start",
        "shall i start",
        "if this looks right",
        "if that looks right",
        "say go ahead",
        "xac nhan",
        "ban xac nhan",
        "neu ban dong y",
        "neu dung toi se bat dau",
        "toi co the bat dau khong",
        "ban muon toi bat dau",
    ]
    launch_markers = ["workflow", "launch", "start", "run", "quy trinh", "bat dau", "chay"]
    return any(marker in folded for marker in confirmation_markers) and any(marker in folded for marker in launch_markers)


def _user_explicitly_confirmed_workflow(messages: List[ChatMessage]) -> bool:
    last_user = ""
    for item in reversed(messages):
        if str(item.role or "").lower() == "user":
            last_user = str(item.content or "").strip()
            break
    if not last_user:
        return False
    folded = _fold_text_for_match(last_user)
    negative_markers = [
        "not yet",
        "do not",
        "dont",
        "don't",
        "stop",
        "wait",
        "hold on",
        "khong",
        "chua",
        "khoan",
        "dung",
    ]
    if any(marker in folded for marker in negative_markers):
        return False
    confirmation_markers = [
        "yes",
        "ok",
        "okay",
        "oke",
        "go ahead",
        "proceed",
        "continue",
        "start",
        "run it",
        "confirm",
        "confirmed",
        "dong y",
        "xac nhan",
        "bat dau",
        "chay di",
        "lam di",
        "tien hanh",
        "trien khai",
        "duoc roi",
    ]
    return any(marker == folded or marker in folded for marker in confirmation_markers)


def _default_confirmation_reply(*, workflow_topic: str, requested_language: str, attachment_records: List[Dict[str, Any]]) -> str:
    attachment_note = " using the attached files as context" if attachment_records else ""
    language_note = f" in {requested_language}" if requested_language else ""
    if str(requested_language or "").strip().lower() == "vietnamese":
        attachment_note_vi = " bằng các tệp đính kèm làm ngữ cảnh" if attachment_records else ""
        return f"Tôi có thể bắt đầu workflow{attachment_note_vi} cho tác vụ này: {workflow_topic}. Xác nhận và tôi sẽ chạy ngay."
    return f"I can start the workflow{attachment_note}{language_note} for this task: {workflow_topic}. Confirm and I will launch it."


def _default_launch_reply(*, workflow_topic: str, requested_language: str, attachment_records: List[Dict[str, Any]]) -> str:
    attachment_note = " using the attached files as context" if attachment_records else ""
    if str(requested_language or "").strip().lower() == "vietnamese":
        attachment_note_vi = " bằng các tệp đính kèm làm ngữ cảnh" if attachment_records else ""
        return f"Đã đủ thông tin. Tôi sẽ khởi chạy workflow{attachment_note_vi} cho tác vụ này ngay bây giờ: {workflow_topic}."
    language_note = f" in {requested_language}" if requested_language else ""
    return f"I have enough information and will launch the workflow{attachment_note}{language_note} now for this task: {workflow_topic}."


def _workflow_request_detected(messages: List[ChatMessage], attachment_records: List[Dict[str, Any]]) -> bool:
    folded = _user_messages_folded(messages)
    if not folded:
        return False
    trigger_tokens = [
        "write",
        "draft",
        "revise",
        "rewrite",
        "edit",
        "improve",
        "analyze",
        "compare",
        "respond to reviewer",
        "viet",
        "chinh sua",
        "sua",
        "phan bien",
        "phan tich",
        "dua vao",
        "search",
        "tim nguon",
        "tim cac nguon",
        "tim tai lieu",
        "viet bai bao",
        "viet bai",
        "viet lai",
        "output tieng viet",
        "file dinh kem",
        "thuat ngu dich thuat",
    ]
    imperative_tokens = ["ban cu", "hay", "giup", "bat dau", "thuc hien", "tien hanh"]
    has_trigger = any(token in folded for token in trigger_tokens)
    has_imperative = any(token in folded for token in imperative_tokens)
    has_attachment_context = bool(attachment_records) and any(
        token in folded
        for token in ["file dinh kem", "tai lieu dinh kem", "hoc thuat ngu", "thuat ngu", "glossary", "spreadsheet", "excel", "docx", "pdf"]
    )
    return (has_trigger or has_attachment_context or (has_imperative and bool(attachment_records))) and (
        len(folded) >= 16 or bool(attachment_records)
    )


def _user_messages_folded(messages: List[ChatMessage]) -> str:
    user_parts = [str(item.content or "").strip() for item in messages if str(item.role or "").lower() == "user" and str(item.content or "").strip()]
    return _fold_text_for_match("\n".join(user_parts))


def _looks_like_new_paper_request(messages: List[ChatMessage], source_manuscript: str) -> bool:
    last_user = ""
    for item in reversed(messages):
        if str(item.role or "").lower() == "user":
            last_user = str(item.content or "").strip()
            break
    if _prompt_requests_revision(last_user) and str(source_manuscript or "").strip():
        return False
    raw_lower = str(last_user or "").lower()
    folded = _user_messages_folded(messages)
    paper_tokens = [
        "write",
        "draft",
        "paper",
        "article",
        "essay",
        "manuscript",
        "viet bai",
        "viet bai bao",
        "bai bao",
        "viet cho toi",
        "soan bai",
        "viết cho tôi",
        "viết bài",
        "viết bài báo",
        "bài báo",
        "bai bao",
    ]
    return any(token in folded for token in paper_tokens) or any(token in raw_lower for token in paper_tokens)


def _conversation_has_topic(messages: List[ChatMessage]) -> bool:
    last_user = ""
    for item in reversed(messages):
        if str(item.role or "").lower() == "user":
            last_user = str(item.content or "").strip()
            break
    raw_lower = str(last_user or "").lower().strip()
    folded = _fold_text_for_match(last_user)
    if not raw_lower and not folded:
        return False
    topic_markers = [
        r"\b(về|ve|about|on|regarding)\b\s+\S+",
        r"\b(topic|chu de|de tai)\b\s*[:\-]?\s*\S+",
    ]
    if any(re.search(pattern, raw_lower, flags=re.IGNORECASE) for pattern in topic_markers):
        return True
    if any(re.search(pattern, folded, flags=re.IGNORECASE) for pattern in topic_markers):
        return True
    generic_terms = {
        "write",
        "draft",
        "paper",
        "article",
        "essay",
        "manuscript",
        "help",
        "please",
        "viet",
        "bai",
        "bao",
        "vietcho",
        "vietchotoi",
        "cho",
        "toi",
        "giup",
        "lam",
        "mot",
    }
    tokens = [token for token in re.findall(r"[a-z0-9]+", folded) if len(token) >= 3 and token not in generic_terms]
    return len(tokens) >= 2


def _conversation_has_word_target(messages: List[ChatMessage]) -> bool:
    return _extract_requested_word_target(messages) is not None


def _extract_requested_word_target(messages: List[ChatMessage]) -> Optional[int]:
    raw_lower = "\n".join(
        str(item.content or "").strip().lower()
        for item in messages
        if str(item.role or "").lower() == "user" and str(item.content or "").strip()
    )
    folded = _user_messages_folded(messages)
    patterns = [
        r"\b(\d{3,5})\s*(?:word|words|tu)\b",
        r"\b(\d{3,5})\s*từ\b",
    ]
    for haystack in (raw_lower, folded):
        for pattern in patterns:
            match = re.search(pattern, haystack, flags=re.IGNORECASE)
            if match:
                value = int(match.group(1))
                if 1000 <= value <= 20000:
                    return value
    return None


def _conversation_has_reference_target(messages: List[ChatMessage]) -> bool:
    return _extract_requested_reference_target(messages) is not None


def _extract_requested_reference_target(messages: List[ChatMessage]) -> Optional[int]:
    raw_lower = "\n".join(
        str(item.content or "").strip().lower()
        for item in messages
        if str(item.role or "").lower() == "user" and str(item.content or "").strip()
    )
    folded = _user_messages_folded(messages)
    patterns = [
        r"\b(\d{1,3})\s*(?:reference|references|ref|refs|source|sources|citation|citations|nguon)\b",
        r"\b(\d{1,3})\s*(?:tài liệu tham khảo|tai lieu tham khao|nguồn|nguon|trích dẫn|trich dan)\b",
    ]
    for haystack in (raw_lower, folded):
        for pattern in patterns:
            match = re.search(pattern, haystack, flags=re.IGNORECASE)
            if match:
                value = int(match.group(1))
                if 4 <= value <= 200:
                    return value
    return None


def _last_user_message_folded(messages: List[ChatMessage]) -> str:
    for item in reversed(messages):
        if str(item.role or "").lower() == "user":
            return _fold_text_for_match(str(item.content or "").strip())
    return ""


def _looks_like_new_paper_request_v2(messages: List[ChatMessage], source_manuscript: str) -> bool:
    last_user = ""
    for item in reversed(messages):
        if str(item.role or "").lower() == "user":
            last_user = str(item.content or "").strip()
            break
    if _prompt_requests_revision(last_user) and str(source_manuscript or "").strip():
        return False
    folded = _user_messages_folded(messages)
    if not folded:
        return False
    paper_tokens = (
        "write",
        "draft",
        "paper",
        "article",
        "essay",
        "manuscript",
        "viet bai",
        "viet bai bao",
        "bai bao",
        "viet cho toi",
        "soan bai",
    )
    return any(token in folded for token in paper_tokens)


def _conversation_has_topic_v2(messages: List[ChatMessage]) -> bool:
    folded = _user_messages_folded(messages)
    if not folded:
        return False
    topic_markers = (
        r"\b(?:ve|about|on|regarding)\b\s+\S+",
        r"\b(?:topic|chu de|de tai)\b\s*[:\-]?\s*\S+",
    )
    if any(re.search(pattern, folded, flags=re.IGNORECASE) for pattern in topic_markers):
        return True
    generic_terms = {
        "write",
        "draft",
        "paper",
        "article",
        "essay",
        "manuscript",
        "help",
        "please",
        "viet",
        "bai",
        "bao",
        "vietcho",
        "vietchotoi",
        "cho",
        "toi",
        "giup",
        "lam",
        "mot",
    }
    tokens = [token for token in re.findall(r"[a-z0-9]+", folded) if len(token) >= 3 and token not in generic_terms]
    return len(tokens) >= 2


def _extract_requested_word_target_v2(messages: List[ChatMessage]) -> Optional[int]:
    folded = _user_messages_folded(messages)
    match = re.search(r"\b(\d{3,5})\s*(?:word|words|tu)\b", folded, flags=re.IGNORECASE)
    if not match:
        return None
    value = int(match.group(1))
    return value if 1000 <= value <= 20000 else None


def _extract_requested_reference_target_v2(messages: List[ChatMessage]) -> Optional[int]:
    folded = _user_messages_folded(messages)
    patterns = (
        r"\b(\d{1,3})\s*(?:reference|references|ref|refs|source|sources|citation|citations|nguon)\b",
        r"\b(\d{1,3})\s*(?:tai lieu tham khao|tai lieu|nguon|trich dan)\b",
    )
    for pattern in patterns:
        match = re.search(pattern, folded, flags=re.IGNORECASE)
        if not match:
            continue
        value = int(match.group(1))
        if 4 <= value <= 200:
            return value
    return None


def _extract_scope_with_llm(
    *,
    cfg: LLMConfig,
    messages: List[ChatMessage],
    source_manuscript: str,
    attachment_context: str,
    known_topic: str,
    known_target_word_count: Optional[int],
    known_reference_target: Optional[int],
) -> Dict[str, Any]:
    try:
        raw = llm_complete_text(
            cfg=cfg,
            prompt=_workflow_scope_extraction_prompt(
                messages=messages,
                source_manuscript=source_manuscript,
                attachment_context=attachment_context,
                known_topic=known_topic,
                known_target_word_count=known_target_word_count,
                known_reference_target=known_reference_target,
            ),
        )
        structured = _extract_json_object(raw)
        if not isinstance(structured, dict):
            return {}
        return structured
    except Exception:
        return {}


def _missing_scope_requirements(
    messages: List[ChatMessage],
    *,
    source_manuscript: str,
    workflow_intent_detected: bool,
) -> List[str]:
    if not workflow_intent_detected:
        return []
    if not _looks_like_new_paper_request_v2(messages, source_manuscript):
        return []
    missing: List[str] = []
    if not _conversation_has_topic_v2(messages):
        missing.append("topic")
    if _extract_requested_word_target_v2(messages) is None:
        missing.append("word_count")
    if _extract_requested_reference_target_v2(messages) is None:
        missing.append("reference_count")
    return missing


def _scope_clarification_reply(requested_language: str, missing_fields: List[str]) -> str:
    needs_topic = "topic" in missing_fields
    needs_words = "word_count" in missing_fields
    needs_refs = "reference_count" in missing_fields
    if str(requested_language or "").strip().lower() == "vietnamese":
        if needs_topic and needs_words and needs_refs:
            return "Trước khi chạy workflow, cho tôi biết rõ đề tài/topic của bài, bài cần khoảng bao nhiêu từ, và bạn muốn khoảng bao nhiêu tài liệu tham khảo/reference."
        if needs_topic and needs_words:
            return "Trước khi chạy workflow, cho tôi biết rõ đề tài/topic của bài và bài cần khoảng bao nhiêu từ."
        if needs_topic and needs_refs:
            return "Trước khi chạy workflow, cho tôi biết rõ đề tài/topic của bài và bạn muốn khoảng bao nhiêu tài liệu tham khảo/reference."
        if needs_topic:
            return "Trước khi chạy workflow, cho tôi biết rõ đề tài/topic cụ thể của bài."
        if needs_words and needs_refs:
            return "Trước khi chạy workflow, cho tôi biết bài cần khoảng bao nhiêu từ và bạn muốn khoảng bao nhiêu tài liệu tham khảo/reference."
        if needs_words:
            return "Trước khi chạy workflow, cho tôi biết bài cần khoảng bao nhiêu từ."
        if needs_refs:
            return "Trước khi chạy workflow, cho tôi biết bạn muốn khoảng bao nhiêu tài liệu tham khảo/reference."
        return "Trước khi chạy workflow, tôi cần thêm một vài thông tin ngắn để chốt yêu cầu."
    if needs_topic and needs_words and needs_refs:
        return "Before I launch the workflow, tell me the exact topic, the target word count, and the approximate number of references you want."
    if needs_topic and needs_words:
        return "Before I launch the workflow, tell me the exact topic and the target word count you want."
    if needs_topic and needs_refs:
        return "Before I launch the workflow, tell me the exact topic and the approximate number of references you want."
    if needs_topic:
        return "Before I launch the workflow, tell me the exact topic you want the paper to cover."
    if needs_words and needs_refs:
        return "Before I launch the workflow, tell me the target word count and the approximate number of references you want."
    if needs_words:
        return "Before I launch the workflow, tell me the target word count you want."
    if needs_refs:
        return "Before I launch the workflow, tell me the approximate number of references you want."
    return "Before I launch the workflow, I need one more short clarification."


def _should_start_workflow_fallback(messages: List[ChatMessage], attachment_records: List[Dict[str, Any]]) -> bool:
    del attachment_records
    return _assistant_requested_workflow_confirmation(messages) and _user_explicitly_confirmed_workflow(messages)


def _chat_response_implies_start(
    *,
    assistant_reply: str,
    readiness: str,
    workflow_topic: str,
    heuristic_should_start: bool,
) -> bool:
    readiness_value = str(readiness or "").strip().lower()
    if readiness_value == "ready_for_workflow":
        return True
    topic_text = str(workflow_topic or "").strip()
    reply = normalize_for_match(str(assistant_reply or "")).lower()
    folded = unicodedata.normalize("NFKD", reply)
    folded = "".join(ch for ch in folded if unicodedata.category(ch) != "Mn")
    folded = folded.replace("đ", "d")
    folded = re.sub(r"\s+", " ", folded).strip()
    launch_markers = [
        "workflow is being launched",
        "launching the workflow",
        "starting the workflow",
        "running the workflow",
        "quy trinh dang duoc khoi chay",
        "quy trinh viet dang duoc khoi chay",
        "toi se bat dau",
        "toi se tien hanh",
        "se bat dau tim kiem tai lieu",
        "bat dau tim kiem tai lieu",
    ]
    return bool(topic_text) and (heuristic_should_start or any(marker in folded for marker in launch_markers))


def _create_job_from_payload(payload: WorkflowRequest) -> str:
    attachment_records = _resolve_attachment_records(payload.attachment_ids)
    job_id = uuid.uuid4().hex
    now = time.time()
    job = JobSnapshot(
        id=job_id,
        status="queued",
        created_at=now,
        updated_at=now,
        request=payload.model_dump(),
        node_statuses=init_statuses(),
        logs=[],
        attachments=[
            {
                "id": item.get("id", ""),
                "filename": item.get("filename", ""),
                "file_type": item.get("file_type", ""),
                "preview": _truncate_chars(item.get("summary") or item.get("extracted_text") or "", 240),
            }
            for item in attachment_records
        ],
    )
    with _jobs_lock:
        _jobs[job_id] = job
    _persist_job(job)
    _job_stop_event(job_id)
    worker = threading.Thread(target=_run_job, args=(job_id, payload), daemon=True)
    worker.start()
    return job_id


def _run_job(job_id: str, payload: WorkflowRequest) -> None:
    statuses = init_statuses()
    _update_job(job_id, status="running", node_statuses=statuses)
    hooks = _make_hooks(job_id, suppress_core_editor=True)

    try:
        _raise_if_stopped(job_id)
        auth_mode = _local_llm_auth_mode()
        if auth_mode != "cli_proxy_oauth":
            raise RuntimeError("Local companion OAuth is required to run the workflow.")
        google_credentials = None
        models, default_model, _source_note = _resolve_models(google_credentials)
        selected_model = payload.model or default_model
        if selected_model not in models:
            models = [selected_model] + models

        attachment_records = _resolve_attachment_records(payload.attachment_ids)
        attachment_context = _attachment_context(attachment_records)
        source_manuscript = _select_source_manuscript(payload.topic, attachment_records)
        research_query = _build_research_query(payload.topic, attachment_records, payload.chat_history)
        requested_output_language = infer_requested_output_language(
            payload.topic,
            "\n".join(
                str(item.content or "").strip()
                for item in payload.chat_history[-12:]
                if str(item.content or "").strip()
            ),
            attachment_context[:4000],
            source_manuscript[:4000],
            fallback=payload.language,
        )
        attachment_summaries = [
            {
                "id": item.get("id", ""),
                "filename": item.get("filename", ""),
                "file_type": item.get("file_type", ""),
                "preview": _truncate_chars(item.get("summary") or item.get("extracted_text") or "", 240),
            }
            for item in attachment_records
        ]
        _update_job(job_id, attachments=attachment_summaries)

        llm_cfg: LLMConfig = build_llm_config(
            "Gemini",
            selected_model,
            models,
            gemini_auth_mode=auth_mode,
            google_credentials=google_credentials,
        )
        _raise_if_stopped(job_id)
        state, logs = execute_workflow(
            topic=payload.topic,
            language=payload.language,
            target_word_count=payload.target_word_count,
            reference_target=payload.reference_target,
            llm_cfg=llm_cfg,
            statuses=statuses,
            hooks=hooks,
            user_prompt=payload.topic,
            chat_history=[item.model_dump() for item in payload.chat_history],
            attachment_context=attachment_context,
            source_manuscript=source_manuscript,
            research_query=research_query,
            search_filters=dict(payload.search_filters or {}),
        )
        manager_guidance = state.get("manager_guidance") if isinstance(state.get("manager_guidance"), dict) else {}
        workflow_language = normalize_language_label(manager_guidance.get("workflow_language"), payload.language)
        output_language = infer_requested_output_language(
            payload.topic,
            "\n".join(
                str(item.content or "").strip()
                for item in payload.chat_history[-12:]
                if str(item.content or "").strip()
            ),
            attachment_context[:4000],
            source_manuscript[:4000],
            str(manager_guidance.get("output_language") or ""),
            fallback=normalize_language_label(manager_guidance.get("output_language"), requested_output_language),
        )
        effective_target_word_count = clamp_int(manager_guidance.get("target_word_count"), 1000, 20000, payload.target_word_count)
        effective_reference_target = clamp_int(manager_guidance.get("reference_target"), 4, 200, payload.reference_target)
        use_translator = bool(manager_guidance.get("use_translator", output_language != workflow_language))
        manager_cfg = llm_config_for_role(llm_cfg, "Manager")
        reviewer_cfg = llm_config_for_role(llm_cfg, "Reviewer")
        writer_cfg = llm_config_for_role(llm_cfg, "Writer")
        editor_cfg = llm_config_for_role(llm_cfg, "Editor")
        translator_cfg = llm_config_for_role(llm_cfg, "Translator")
        logs = _filter_online_logs(logs)
        statuses["Reviewer"] = statuses.get("Reviewer", "Pending")
        statuses["Editor"] = "Pending"
        statuses["Translator"] = "Pending"
        _update_job(job_id, node_statuses=dict(statuses), logs=list(logs))

        def add_online_log(message: str) -> None:
            timestamp = time.strftime("%H:%M:%S")
            line = f"[{timestamp}] {message}"
            logs.append(line)
            _append_log(job_id, line)

        abstract_target = min(250, max(120, int(effective_target_word_count * 0.12)))
        online_body_target = max(320, effective_target_word_count - estimated_reference_words(effective_reference_target) - abstract_target)
        draft_body = _strip_references_section(state.get("final_draft") or "")
        if not draft_body:
            draft_body = _strip_references_section(state.get("final_formatted") or "")

        if draft_body:
            _raise_if_stopped(job_id)
            statuses["Reviewer"] = "Processing"
            _update_job(job_id, node_statuses=dict(statuses))
            add_online_log(
                f"Online Reviewer started: critiquing draft for argument quality, target length ~{online_body_target} words, and citation limit {effective_reference_target}."
            )
            review_parts: List[str] = []
            for delta in llm_stream_text(
                cfg=reviewer_cfg,
                prompt=_online_reviewer_feedback_prompt(
                    language=workflow_language,
                    draft=draft_body,
                    target_word_count=online_body_target,
                    reference_target=effective_reference_target,
                    manager_brief=str(manager_guidance.get("reviewer_brief", "") or ""),
                ),
                should_stop=lambda: _raise_if_stopped(job_id),
            ):
                _raise_if_stopped(job_id)
                review_parts.append(delta)
                _update_job(job_id, review_feedback="".join(review_parts).strip())
            review_feedback = "".join(review_parts).strip()
            statuses["Reviewer"] = "Done"
            _update_job(job_id, node_statuses=dict(statuses))
            add_online_log("Online Reviewer done: critique generated for rewrite.")

            _raise_if_stopped(job_id)
            statuses["Writer"] = "Processing"
            _update_job(job_id, node_statuses=dict(statuses))
            add_online_log("Online Writer rewrite started: revising draft from reviewer feedback.")
            rewritten_body = llm_complete_text(
                cfg=writer_cfg,
                prompt=_online_rewrite_prompt(
                    language=workflow_language,
                    draft=draft_body,
                    target_word_count=online_body_target,
                    review_feedback=review_feedback,
                    reference_target=effective_reference_target,
                    manager_brief=str(manager_guidance.get("writer_brief", "") or ""),
                ),
                should_stop=lambda: _raise_if_stopped(job_id),
            ).strip()
            if rewritten_body:
                draft_body = rewritten_body
                state["reviewed_draft"] = rewritten_body
                _update_job(job_id, draft=rewritten_body)
            statuses["Writer"] = "Done"
            _update_job(job_id, node_statuses=dict(statuses))
            add_online_log("Online Writer rewrite done: revised manuscript body prepared.")
        else:
            statuses["Reviewer"] = "Done"
            statuses["Writer"] = "Done"
            _update_job(job_id, node_statuses=dict(statuses))
            draft_body = _strip_references_section(state.get("final_formatted") or "")

        _raise_if_stopped(job_id)
        add_online_log("Online Abstract synthesis started: summarizing full manuscript body.")
        editor_language = workflow_language if use_translator else output_language
        abstract_block = llm_complete_text(
            cfg=editor_cfg,
            prompt=_online_abstract_prompt(
                language=editor_language,
                manuscript_body=draft_body,
                target_word_count=effective_target_word_count,
                manager_brief=str(manager_guidance.get("editor_brief", "") or manager_guidance.get("writer_brief", "") or ""),
                source_language=workflow_language,
            ),
            should_stop=lambda: _raise_if_stopped(job_id),
        ).strip()
        manuscript_with_abstract = f"{abstract_block}\n\n{draft_body}".strip() if abstract_block else draft_body
        _update_job(job_id, draft=manuscript_with_abstract)
        add_online_log("Online Abstract synthesis done: abstract added to draft.")

        _raise_if_stopped(job_id)
        statuses["Editor"] = "Processing"
        _update_job(job_id, node_statuses=dict(statuses))
        add_online_log("Online Editor started: polishing abstract, manuscript, and rebuilding references.")
        polished_manuscript = llm_complete_text(
            cfg=editor_cfg,
            prompt=_online_editor_prompt(
                language=editor_language,
                draft=manuscript_with_abstract,
                target_word_count=effective_target_word_count - estimated_reference_words(effective_reference_target),
                manager_brief=str(manager_guidance.get("editor_brief", "") or ""),
                source_language=editor_language,
            ),
            should_stop=lambda: _raise_if_stopped(job_id),
        ).strip()
        final_body_for_output = polished_manuscript or manuscript_with_abstract

        reference_match_text = polished_manuscript or manuscript_with_abstract
        statuses["Editor"] = "Done"
        _update_job(job_id, node_statuses=dict(statuses))
        add_online_log("Online Editor done: source-language manuscript polished.")

        if use_translator and final_body_for_output:
            _raise_if_stopped(job_id)
            statuses["Translator"] = "Processing"
            _update_job(job_id, node_statuses=dict(statuses))
            add_online_log("Online Translator started: translating manuscript with attachment-aware terminology guidance.")
            translated_manuscript = llm_complete_text(
                cfg=translator_cfg,
                prompt=_online_translator_prompt(
                    language=output_language,
                    draft=final_body_for_output,
                    source_language=workflow_language,
                    manager_brief=str(manager_guidance.get("translator_brief", "") or manager_guidance.get("editor_brief", "") or ""),
                    attachment_context=attachment_context,
                ),
                should_stop=lambda: _raise_if_stopped(job_id),
            ).strip()
            if translated_manuscript:
                final_body_for_output = translated_manuscript
            statuses["Translator"] = "Done"
            _update_job(job_id, node_statuses=dict(statuses))
            add_online_log("Online Translator done: final manuscript translated.")
        else:
            statuses["Translator"] = "Done"
            _update_job(job_id, node_statuses=dict(statuses))

        detected_output_language = guess_language_label(final_body_for_output[:4000])
        if final_body_for_output and detected_output_language != output_language:
            add_online_log(
                f"Online Translator safeguard started: detected {detected_output_language}, forcing final manuscript to {output_language}."
            )
            final_body_for_output = _ensure_output_language(
                cfg=translator_cfg if use_translator else manager_cfg,
                draft=final_body_for_output,
                target_language=output_language,
                should_stop=lambda: _raise_if_stopped(job_id),
            )
            add_online_log(
                f"Online Translator safeguard done: detected={guess_language_label(final_body_for_output[:4000])}, requested={output_language}."
            )

        localized_manuscript = localize_in_text_citations(final_body_for_output, output_language).strip()
        if localized_manuscript:
            _update_job(job_id, final_markdown=localized_manuscript)

        references_section = build_reference_section(
            state.get("metadata_list", []),
            text=reference_match_text,
            limit=effective_reference_target,
        )
        final_markdown = localized_manuscript
        if references_section:
            final_markdown = f"{final_markdown}\n\n{references_section}".strip()
        add_online_log("Quality Check started: verifying structure, word count, references, and output language.")
        qa_summary = _build_quality_summary(
            final_markdown=final_markdown,
            effective_target_word_count=effective_target_word_count,
            effective_reference_target=effective_reference_target,
            output_language=output_language,
            task_mode=str(manager_guidance.get("task_mode", "") or ""),
        )
        add_online_log(
            f"Online Finalization done: final manuscript ready with {effective_reference_target} reference slots and {count_words(final_markdown)} words."
        )
        add_online_log("Quality Check done: final manuscript package is ready for Markdown and Word export.")
        _update_job(
            job_id,
            status="completed",
            logs=list(logs),
            node_statuses=dict(statuses),
            outline=state.get("plan", ""),
            papers=state.get("metadata_list", []),
            review_feedback=review_feedback if "review_feedback" in locals() else "",
            qa_summary=qa_summary,
            attachments=attachment_summaries,
            draft=manuscript_with_abstract or state.get("reviewed_draft") or state.get("final_draft", ""),
            final_markdown=final_markdown,
            actual_word_count=count_words(final_markdown),
        )
        _clear_job_stop(job_id)
    except JobCancelledError as exc:
        for node in NODE_ORDER:
            if statuses.get(node) == "Processing":
                statuses[node] = "Error"
        _update_job(
            job_id,
            status="cancelled",
            node_statuses=dict(statuses),
            error=str(exc),
        )
        _append_log(job_id, "Workflow cancelled by user.")
        _clear_job_stop(job_id)
    except Exception as exc:  # noqa: BLE001
        for node in NODE_ORDER:
            if statuses.get(node) == "Processing":
                statuses[node] = "Error"
        message = f"{exc}"
        trace = traceback.format_exc(limit=8)
        _update_job(
            job_id,
            status="error",
            node_statuses=dict(statuses),
            error=message,
        )
        _append_log(job_id, f"Workflow failed: {message}")
        _append_log(job_id, trace)
        _clear_job_stop(job_id)


@app.get("/api/health")
def health() -> Dict[str, Any]:
    is_local_companion = _is_local_companion_runtime()
    _client_id, _client_secret, _project_id, oauth_client_source = _effective_google_oauth_client_config_values()
    local_auth_entries = _load_local_cli_proxy_auth_entries() if is_local_companion else []
    local_cli_proxy = bool(is_local_companion and (cliproxy_available() or local_auth_entries))
    cli_proxy_owner = os.getenv("RESEARCH_COMPANION_CLI_PROXY_OWNED", "").strip().lower() in {"1", "true", "yes", "on"}
    auth_state = _read_local_cli_proxy_auth_state() if local_cli_proxy else {"status": "unknown", "reason": ""}
    auth_mode = _local_llm_auth_mode() if local_cli_proxy else "local_companion_required"
    return {
        "ok": True,
        "oauth_store_backend": OAUTH_STORE_BACKEND,
        "default_provider": DEFAULT_LLM_PROVIDER,
        "default_model": DEFAULT_GEMINI_MODEL,
        "public_backend_url": PUBLIC_BACKEND_URL,
        "local_companion_mode": is_local_companion,
        "google_oauth_client_source": oauth_client_source,
        "cli_proxy_api_available": local_cli_proxy,
        "cli_proxy_config_owner": "research_companion" if cli_proxy_owner else "",
        "cli_proxy_auth_dir": str(_local_cli_proxy_auth_dir()),
        "llm_auth_mode": auth_mode,
        "llm_auth_reason": str(auth_state.get("reason") or "").strip(),
    }


@app.get("/api/quota")
def quota() -> Dict[str, Any]:
    if not _is_local_companion_runtime():
        return {"available": False, "reason": "Quota is only available through the local companion runtime."}
    if not _load_local_cli_proxy_auth_entries():
        return {"available": False, "reason": "Start ResearchCompanion.exe and sign in through the local companion first."}
    try:
        payload = _build_local_quota_payload()
        if payload.get("available"):
            _clear_local_cli_proxy_auth_state()
        elif _is_local_cli_proxy_auth_error_message(str(payload.get("reason") or "")):
            _mark_local_cli_proxy_auth_state("invalid", str(payload.get("reason") or ""))
        return payload
    except HTTPException as exc:
        detail = str(exc.detail or exc.status_code)
        if _is_local_cli_proxy_auth_error_message(detail):
            _mark_local_cli_proxy_auth_state("invalid", detail)
        return {
            "available": False,
            "reason": detail,
        }
    except Exception as exc:
        detail = str(exc)
        if _is_local_cli_proxy_auth_error_message(detail):
            _mark_local_cli_proxy_auth_state("invalid", detail)
        return {
            "available": False,
            "reason": detail,
        }


@app.post("/api/resource-search")
def resource_search(payload: ResourceSearchRequest) -> Dict[str, Any]:
    source = str(payload.source or "").strip().lower()
    topic = str(payload.topic or "").strip()
    max_results = clamp_int(payload.max_results, 1, 100, 25)
    if not topic:
        return {"items": []}

    try:
        if source == "scopus":
            if not SCOPUS_API_KEY.strip():
                raise HTTPException(status_code=503, detail="Scopus API key is not configured on the cloud backend.")
            items = scopus_search(api_key=SCOPUS_API_KEY, topic=topic, max_results=max_results)
        elif source == "core":
            if not CORE_API_KEY.strip():
                raise HTTPException(status_code=503, detail="CORE API key is not configured on the cloud backend.")
            items = core_search(api_key=CORE_API_KEY, topic=topic, max_results=max_results)
        elif source == "openalex":
            items = openalex_search(topic=topic, max_results=max_results)
        elif source == "arxiv":
            items = arxiv_search(topic=topic, max_results=max_results)
        else:
            raise HTTPException(status_code=400, detail="Unsupported resource source.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"{source} resource search failed: {exc}") from exc

    return {"items": items}


@app.post("/api/resource-unpaywall")
def resource_unpaywall(payload: ResourceUnpaywallRequest) -> Dict[str, Any]:
    doi = str(payload.doi or "").strip()
    if not doi:
        return {"item": {}}
    if not UNPAYWALL_EMAIL.strip():
        raise HTTPException(status_code=503, detail="Unpaywall email is not configured on the cloud backend.")
    try:
        item = unpaywall_lookup(doi=doi, email=UNPAYWALL_EMAIL)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Unpaywall lookup failed: {exc}") from exc
    return {"item": item}


@app.get("/")
def index() -> FileResponse:
    return FileResponse(os.path.join(WEB_DIR, "index.html"))


@app.get("/api/oauth/google/start")
def oauth_google_start() -> Dict[str, Any]:
    _ensure_oauth_store_ready()
    flow_id = uuid.uuid4().hex
    state_token = uuid.uuid4().hex
    auth_url, client_id = _google_auth_url_for_flow(flow_id, state_token)
    now = time.time()
    _upsert_oauth_flow(
        flow_id,
        {
            "flow_id": flow_id,
            "state": state_token,
            "status": "pending",
            "provider": "google_oauth",
            "client_id": client_id,
            "auth_url": auth_url,
            "created_at": now,
            "updated_at": now,
        },
    )
    return {"flow_id": flow_id, "auth_url": auth_url, "status": "pending"}


@app.get("/api/oauth/google/status")
def oauth_google_status(flow_id: str) -> Dict[str, Any]:
    _ensure_oauth_store_ready()
    entry = _oauth_flow(flow_id=flow_id)
    if not entry:
        raise HTTPException(status_code=404, detail="OAuth flow not found")
    return {"flow": entry}


@app.get("/api/oauth/google/session")
def oauth_google_session(oauth_session_id: str) -> Dict[str, Any]:
    _ensure_oauth_store_ready()
    entry = _google_session_entry(oauth_session_id)
    if not entry:
        raise HTTPException(status_code=404, detail="OAuth session not found")
    auth_file = str(entry.get("auth_file") or "").strip()
    credentials = _session_credentials_from_entry(oauth_session_id, entry)
    if credentials is None:
        _delete_google_session(oauth_session_id)
        raise HTTPException(status_code=404, detail="OAuth session expired")
    entry["updated_at"] = time.time()
    _upsert_google_session(oauth_session_id, entry)
    return {
        "connected": True,
        "oauth_session_id": oauth_session_id,
        "email": str(entry.get("email") or "").strip(),
        "auth_file": auth_file,
    }


class GoogleOAuthLogoutRequest(BaseModel):
    oauth_session_id: str = ""


@app.post("/api/oauth/google/logout")
def oauth_google_logout(payload: GoogleOAuthLogoutRequest) -> Dict[str, Any]:
    _ensure_oauth_store_ready()
    _delete_google_session(payload.oauth_session_id.strip())
    return {"ok": True}


@app.get("/oauth/google/callback")
def oauth_google_callback(
    state: str = "",
    code: str = "",
    error: str = "",
) -> HTMLResponse:
    _ensure_oauth_store_ready()
    flow_entry = _oauth_flow(state_token=state)
    matched_flow_id = str((flow_entry or {}).get("flow_id") or "").strip()

    if error:
        if matched_flow_id:
            _update_oauth_flow(matched_flow_id, status="error", error=f"OAuth provider returned error: {error}")
        return HTMLResponse(
            _oauth_success_html(f"OAuth provider returned error: {error}", success=False),
            status_code=400,
        )

    if not matched_flow_id or not code:
        return HTMLResponse(
            _oauth_success_html("Missing authorization code or state.", success=False),
            status_code=400,
        )

    client_section = _google_client_section()
    token_uri = str(client_section.get("token_uri") or "https://oauth2.googleapis.com/token").strip()
    client_id = str(client_section.get("client_id") or "").strip()
    client_secret = str(client_section.get("client_secret") or "").strip()
    try:
        response = requests.post(
            token_uri,
            data={
                "code": code,
                "client_id": client_id,
                "client_secret": client_secret,
                "redirect_uri": _google_redirect_uri(),
                "grant_type": "authorization_code",
            },
            timeout=30,
        )
        response.raise_for_status()
        token_info = response.json()
        access_token = str(token_info.get("access_token") or "").strip()
        if not access_token:
            raise RuntimeError("Google token response did not include an access token.")
        if str(token_info.get("refresh_token") or "").strip():
            token_info["client_id"] = client_id
            token_info["client_secret"] = client_secret
        userinfo = google_oauth_userinfo(access_token)
        email = str(userinfo.get("email") or "").strip() or "google_user"
        auth_file = google_auth_file_save(token_info, email)
        session_id = uuid.uuid4().hex
        _upsert_google_session(
            session_id,
            {
                "oauth_session_id": session_id,
                "email": email,
                "auth_file": auth_file,
                "token_info": token_info,
                "created_at": time.time(),
                "updated_at": time.time(),
            },
        )
        _update_oauth_flow(
            matched_flow_id,
            status="done",
            email=email,
            auth_file=auth_file,
            oauth_session_id=session_id,
        )
        return HTMLResponse(
            _oauth_success_html(f"Connected Google OAuth for {email}.", success=True),
            status_code=200,
        )
    except Exception as exc:
        _update_oauth_flow(matched_flow_id, status="error", error=str(exc))
        return HTMLResponse(
            _oauth_success_html(str(exc), success=False),
            status_code=400,
        )


@app.post("/api/uploads")
async def upload_attachments(
    files: List[UploadFile] = File(...),
) -> Dict[str, Any]:
    _ensure_oauth_store_ready()
    credentials = _runtime_google_credentials()
    items: List[Dict[str, Any]] = []
    for file in files:
        filename = str(file.filename or "upload").strip() or "upload"
        mime_type = str(file.content_type or _guess_mime_type(filename)).strip() or "application/octet-stream"
        data = await file.read()
        if not data:
            continue
        if len(data) > 25 * 1024 * 1024:
            raise HTTPException(status_code=400, detail=f"{filename}: file is larger than 25 MB.")

        upload_id = uuid.uuid4().hex
        blob_path = _upload_blob_path(upload_id, filename)
        Path(blob_path).write_bytes(data)

        extracted_text = _extract_attachment_text(filename, mime_type, data)
        summary = ""
        if not extracted_text and mime_type.startswith("image/"):
            summary = _summarize_image_bytes(
                image_bytes=data,
                mime_type=mime_type,
                filename=filename,
                credentials=credentials,
            )

        extension = Path(filename).suffix.lower()
        if mime_type.startswith("image/"):
            file_type = "image"
        elif extension in {".pdf"}:
            file_type = "pdf"
        elif extension in {".docx"}:
            file_type = "word"
        elif extension in {".xlsx", ".xlsm", ".csv", ".tsv"}:
            file_type = "spreadsheet"
        else:
            file_type = "document"

        record = {
            "id": upload_id,
            "filename": filename,
            "mime_type": mime_type,
            "file_type": file_type,
            "size_bytes": len(data),
            "created_at": time.time(),
            "path": blob_path,
            "extracted_text": _truncate_chars(extracted_text, 24000),
            "summary": _truncate_chars(summary, 8000),
        }
        _write_upload_record(upload_id, record)
        items.append(
            {
                "id": upload_id,
                "filename": filename,
                "mime_type": mime_type,
                "file_type": file_type,
                "size_bytes": len(data),
                "has_text": bool(record["extracted_text"] or record["summary"]),
                "preview": _truncate_chars(record["summary"] or record["extracted_text"], 260),
            }
        )
    return {"items": items}


@app.post("/api/chat-turn")
def chat_turn(payload: ChatTurnRequest) -> Dict[str, Any]:
    _ensure_oauth_store_ready()
    _require_local_companion_oauth()
    auth_mode = _local_llm_auth_mode()
    google_credentials = None

    models, default_model, _source_note = _resolve_models(google_credentials)
    selected_model = payload.model or default_model
    if selected_model not in models:
        models = [selected_model] + models

    llm_cfg: LLMConfig = build_llm_config(
        "Gemini",
        selected_model,
        models,
        gemini_auth_mode=auth_mode,
        google_credentials=google_credentials,
    )
    manager_cfg = llm_config_for_role(llm_cfg, "WorkflowChat")

    attachment_records = _resolve_attachment_records(payload.attachment_ids)
    attachment_context = _attachment_context(attachment_records)
    last_user_prompt = ""
    for item in reversed(payload.messages):
        if str(item.role or "user").lower() == "user":
            last_user_prompt = str(item.content or "").strip()
            break
    source_manuscript = _select_source_manuscript(last_user_prompt, attachment_records)
    heuristic_query = _build_research_query(last_user_prompt, attachment_records, payload.messages)
    requested_language = infer_requested_output_language(
        last_user_prompt,
        _chat_messages_block(payload.messages),
        attachment_context[:4000],
        source_manuscript[:4000],
        fallback=payload.language,
    )

    raw = llm_complete_text(
        cfg=manager_cfg,
        prompt=_workflow_chat_protocol_prompt(
            language=requested_language,
            target_word_count=payload.target_word_count,
            reference_target=payload.reference_target,
            messages=payload.messages,
            attachment_context=attachment_context,
            source_manuscript=source_manuscript,
            heuristic_research_query=heuristic_query,
        ),
    )
    structured = _extract_json_object(raw)

    workflow_intent_detected = structured.get("workflow_intent_detected") is True
    task_mode = str(structured.get("task_mode") or "").strip().lower()
    extracted_topic = str(structured.get("topic") or "").strip()
    raw_word_target = structured.get("target_word_count")
    raw_reference_target = structured.get("reference_target")
    extracted_target_word_count = raw_word_target if isinstance(raw_word_target, int) and 1000 <= raw_word_target <= 20000 else None
    extracted_reference_target = raw_reference_target if isinstance(raw_reference_target, int) and 4 <= raw_reference_target <= 200 else None
    raw_missing = structured.get("missing_requirements")
    missing_scope_requirements = []
    if isinstance(raw_missing, list):
        for item in raw_missing:
            value = str(item or "").strip().lower()
            if value in {"topic", "word_count", "reference_count"} and value not in missing_scope_requirements:
                missing_scope_requirements.append(value)

    assistant_reply = str(structured.get("assistant_reply") or "").strip()
    workflow_topic = str(
        structured.get("workflow_topic")
        or extracted_topic
        or last_user_prompt
        or "Use the conversation and attached files to continue the workflow."
    ).strip()
    should_start = structured.get("should_start_workflow")
    if not isinstance(should_start, bool):
        should_start = False
    readiness = str(structured.get("readiness") or ("ready_for_workflow" if should_start else "chat_only")).strip()
    reason = str(structured.get("reason") or "").strip()
    if should_start:
        readiness = "ready_for_workflow"
    elif readiness == "ready_for_workflow":
        readiness = "needs_clarification"

    auto_ready_for_workflow = (
        workflow_intent_detected
        and task_mode in {"new_paper", "revise_manuscript", "critique_and_rewrite", "summarize", "compare", "literature_review"}
        and bool((workflow_topic or extracted_topic or "").strip())
        and not missing_scope_requirements
    )
    if task_mode == "new_paper":
        auto_ready_for_workflow = auto_ready_for_workflow and extracted_target_word_count is not None and extracted_reference_target is not None
    if auto_ready_for_workflow:
        should_start = True
        readiness = "ready_for_workflow"
        assistant_reply_lower = assistant_reply.lower()
        if (
            not assistant_reply
            or "confirm" in assistant_reply_lower
            or "xác nhận" in assistant_reply_lower
            or "bạn có muốn" in assistant_reply_lower
            or "do you want me to start" in assistant_reply_lower
            or "i can start the workflow" in assistant_reply_lower
            or "tôi có thể bắt đầu workflow" in assistant_reply_lower
        ):
            assistant_reply = _default_launch_reply(
                workflow_topic=workflow_topic,
                requested_language=requested_language,
                attachment_records=attachment_records,
            )

    if missing_scope_requirements:
        should_start = False
        readiness = "needs_clarification"
        assistant_reply = assistant_reply or _scope_clarification_reply(requested_language, missing_scope_requirements)
        reason = reason or "Waiting for the user to fill in the missing workflow requirements."
    elif not assistant_reply:
        assistant_reply = (
            "I have enough source material and will launch the workflow now. I will route the task, gather evidence, draft the manuscript, run review, perform quality checks, and prepare exportable output."
            if should_start
            else "I understand. Continue the conversation or give a concrete task when you want me to start the workflow."
        )

    workflow_target_word_count = extracted_target_word_count
    workflow_reference_target = extracted_reference_target
    if workflow_target_word_count is None or workflow_reference_target is None:
        if task_mode in {"revise_manuscript", "critique_and_rewrite"} and str(source_manuscript or "").strip():
            inferred_word_count, inferred_reference_target = infer_revision_targets(
                user_prompt=last_user_prompt,
                source_manuscript=source_manuscript,
                fallback_word_count=4000,
                fallback_reference_target=24,
            )
            workflow_target_word_count = workflow_target_word_count or inferred_word_count
            workflow_reference_target = workflow_reference_target or inferred_reference_target
        else:
            workflow_target_word_count = workflow_target_word_count or clamp_int(payload.target_word_count, 1000, 20000, 4000)
            workflow_reference_target = workflow_reference_target or clamp_int(payload.reference_target, 4, 200, 24)

    job_id = ""
    if should_start:
        workflow_payload = WorkflowRequest(
            topic=workflow_topic,
            language=requested_language,
            provider=payload.provider,
            model=payload.model,
            target_word_count=workflow_target_word_count,
            reference_target=workflow_reference_target,
            chat_history=payload.messages[-16:],
            attachment_ids=payload.attachment_ids,
            search_filters=payload.search_filters,
        )
        job_id = _create_job_from_payload(workflow_payload)

    return {
        "assistant_reply": assistant_reply,
        "should_start_workflow": should_start,
        "workflow_topic": workflow_topic,
        "reason": reason,
        "readiness": readiness,
        "job_id": job_id,
    }


@app.post("/api/jobs")
def create_job(payload: WorkflowRequest) -> Dict[str, Any]:
    _ensure_oauth_store_ready()
    _require_local_companion_oauth()
    attachment_records = _resolve_attachment_records(payload.attachment_ids)
    attachment_context = _attachment_context(attachment_records)
    source_manuscript = _select_source_manuscript(payload.topic, attachment_records)
    inferred_language = infer_requested_output_language(
        payload.topic,
        attachment_context[:4000],
        source_manuscript[:4000],
        fallback=payload.language,
    )
    target_word_count = payload.target_word_count
    reference_target = payload.reference_target
    if target_word_count is None or reference_target is None:
        if _prompt_requests_revision(payload.topic) and str(source_manuscript or "").strip():
            inferred_word_count, inferred_reference_target = infer_revision_targets(
                user_prompt=payload.topic,
                source_manuscript=source_manuscript,
                fallback_word_count=4000,
                fallback_reference_target=24,
            )
            target_word_count = target_word_count or inferred_word_count
            reference_target = reference_target or inferred_reference_target
        else:
            raise HTTPException(
                status_code=400,
                detail="target_word_count and reference_target are required before starting a new paper workflow.",
            )
    payload = payload.model_copy(
        update={
            "language": inferred_language,
            "target_word_count": target_word_count,
            "reference_target": reference_target,
        }
    )
    job_id = _create_job_from_payload(payload)
    return {"job_id": job_id}


@app.post("/api/jobs/{job_id}/stop")
def stop_job(job_id: str) -> Dict[str, Any]:
    job = _job_entry(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.status in {"completed", "error", "cancelled"}:
        return {"ok": True, "status": job.status}
    _update_job(job_id, status="cancelling")
    _job_stop_event(job_id).set()
    _append_log(job_id, f"[{time.strftime('%H:%M:%S')}] Stop requested by user.")
    return {"ok": True, "status": "cancelling"}


@app.get("/api/jobs/{job_id}")
def get_job(job_id: str) -> Dict[str, Any]:
    job = _job_entry(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    return _sanitize_job_payload(_job_to_dict(job))


@app.get("/api/jobs/{job_id}/docx")
def download_job_docx(job_id: str) -> Response:
    job = _job_entry(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    if str(job.status or "").strip().lower() != "completed" or not str(job.final_markdown or "").strip():
        raise HTTPException(status_code=400, detail="Final manuscript is not ready yet.")

    docx_bytes = _markdown_to_docx_bytes(job.final_markdown)
    filename = f"research-workflow-{job_id[:8]}.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/api/render/docx")
def render_docx(payload: DocxRenderRequest) -> Response:
    markdown = str(payload.markdown or "").strip()
    if not markdown:
        raise HTTPException(status_code=400, detail="Markdown content is required.")
    docx_bytes = _markdown_to_docx_bytes(markdown)
    filename = _safe_docx_filename(payload.filename)
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.delete("/api/jobs/{job_id}")
def delete_job(job_id: str) -> Dict[str, Any]:
    existing = _job_entry(job_id)
    if existing is None:
        raise HTTPException(status_code=404, detail="Job not found")
    with _jobs_lock:
        if job_id in _jobs:
            del _jobs[job_id]
    _clear_job_stop(job_id)
    _delete_job_entry(job_id)
    return {"deleted": True}


if os.path.isdir(WEB_DIR):
    app.mount("/", StaticFiles(directory=WEB_DIR, html=True), name="web")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "backend_api:app",
        host="0.0.0.0",
        port=int(os.getenv("PORT", "8080")),
        reload=False,
    )
